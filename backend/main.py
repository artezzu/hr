from datetime import datetime, timedelta
from typing import Annotated, List, Optional, Dict, Set, Any, Union
from fastapi import FastAPI, Depends, HTTPException, status, WebSocket, WebSocketDisconnect, File, UploadFile, Form, BackgroundTasks, Request, Query, Body
from fastapi import Path as FastAPIPath  # Явное переименование для избежания конфликта
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.middleware.cors import CORSMiddleware
from jose import JWTError, jwt
from passlib.context import CryptContext
from sqlalchemy.orm import Session, joinedload, contains_eager
import models, schemas
from database import engine, get_db, SessionLocal
from bot import send_message_to_candidate
from hh_api import HeadHunterParser
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
import os
import json
import logging
# Импортируем Path из pathlib с явным псевдонимом
from pathlib import Path as PathLib
import PyPDF2
import docx
import tempfile
import re
import shutil
from docx import Document
from config import (
    SECRET_KEY, ALGORITHM, ACCESS_TOKEN_EXPIRE_MINUTES,
    RESUME_DIR, CORS_ORIGINS, BASE_DIR
)
from werkzeug.utils import secure_filename
import asyncio
import hashlib
import random
import string
import zipfile
import calendar
from collections import defaultdict
import traceback
import time
from sqlalchemy import func, or_, and_, desc, distinct
import base64
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import smtplib
from dotenv import load_dotenv
from auth import create_access_token, get_current_user
from email_validator import validate_email, EmailNotValidError
from dateutils import date_to_datetime
from uuid import uuid4

models.Base.metadata.create_all(bind=engine)

app = FastAPI(title="Abstract HR Platform API")

# Security config
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,  # Используем настройки из config.py
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# WebSocket connections manager
class ConnectionManager:
    def __init__(self):
        self.active_connections: List[WebSocket] = []
        self.authenticated_users: Dict[str, WebSocket] = {}
    
    async def connect(self, websocket: WebSocket, user_email: str):
        # Обратите внимание, что websocket.accept() вызывается до этого метода
        self.active_connections.append(websocket)
        self.authenticated_users[user_email] = websocket
        print(f"WebSocket connected for user {user_email}. Total connections: {len(self.active_connections)}")

    def disconnect(self, websocket: WebSocket):
        if websocket in self.active_connections:
            self.active_connections.remove(websocket)
            
            # Также удаляем из словаря аутентифицированных пользователей
            for email, conn in list(self.authenticated_users.items()):
                if conn == websocket:
                    del self.authenticated_users[email]
                    print(f"User {email} disconnected")
                    break
                    
            print(f"WebSocket disconnected. Remaining connections: {len(self.active_connections)}")
        else:
            print("Attempted to disconnect a WebSocket that was not in active connections")

    async def broadcast(self, message: dict):
        if not self.active_connections:
            print("No active connections to broadcast to")
            return
            
        print(f"Broadcasting message to {len(self.active_connections)} clients: {message.get('type', 'no-type')}")
        disconnected = []
        
        for connection in self.active_connections:
            try:
                await connection.send_json(message)
            except Exception as e:
                print(f"Error broadcasting message: {e}")
                disconnected.append(connection)
        
        # Отключаем соединения с ошибками после итерации
        for conn in disconnected:
            self.disconnect(conn)

# Создаем экземпляр ConnectionManager
manager = ConnectionManager()

def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def get_user(db: Session, email: str):
    return db.query(models.User).filter(models.User.email == email).first()

def authenticate_user(db: Session, email: str, password: str):
    user = get_user(db, email)
    print(f"Attempting to authenticate user: {email}")
    print(f"User found: {user is not None}")
    if user:
        print(f"Verifying password...")
        is_valid = verify_password(password, user.hashed_password)
        print(f"Password valid: {is_valid}")
        if is_valid:
            return user
    return False

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(
    token: Annotated[str, Depends(oauth2_scheme)],
    db: Session = Depends(get_db)
):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise credentials_exception
        token_data = schemas.TokenData(email=email)
    except JWTError:
        raise credentials_exception
    user = get_user(db, email=token_data.email)
    if user is None:
        raise credentials_exception
    return user

@app.post("/token", response_model=schemas.Token)
async def login_for_access_token(
    form_data: Annotated[OAuth2PasswordRequestForm, Depends()],
    db: Session = Depends(get_db)
):
    user = authenticate_user(db, form_data.username, form_data.password)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect email or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.email}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/users/", response_model=schemas.User)
def create_user(user: schemas.UserCreate, db: Session = Depends(get_db)):
    db_user = get_user(db, email=user.email)
    if db_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    hashed_password = get_password_hash(user.password)
    db_user = models.User(
        email=user.email,
        full_name=user.full_name,
        hashed_password=hashed_password,
        role=user.role,
        position=user.position
    )
    db.add(db_user)
    db.commit()
    db.refresh(db_user)
    return db_user

@app.get("/users/me/", response_model=schemas.User)
async def read_users_me(
    current_user: Annotated[schemas.User, Depends(get_current_user)]
):
    return current_user

@app.put("/users/me/update", response_model=schemas.User)
async def update_current_user(
    user_update: schemas.UserUpdate,
    current_user: Annotated[schemas.User, Depends(get_current_user)],
    db: Session = Depends(get_db)
):
    try:
        # Получаем пользователя из базы данных
        db_user = db.query(models.User).filter(models.User.id == current_user.id).first()
        if not db_user:
            raise HTTPException(status_code=404, detail="Пользователь не найден")
        
        # Проверяем что имя не пустое
        if not user_update.full_name or not user_update.full_name.strip():
            raise HTTPException(status_code=400, detail="Имя не может быть пустым")
        
        # Обновляем данные пользователя
        db_user.full_name = user_update.full_name
        
        if user_update.email is not None:
            # Проверяем, что email не занят другим пользователем
            existing_user = db.query(models.User).filter(
                models.User.email == user_update.email, 
                models.User.id != current_user.id
            ).first()
            if existing_user:
                raise HTTPException(status_code=400, detail="Email уже используется другим пользователем")
            db_user.email = user_update.email
        
        db.commit()
        db.refresh(db_user)
        return db_user
    except HTTPException:
        # Пробрасываем уже сформированные HTTP исключения
        raise
    except Exception as e:
        # Логируем непредвиденные ошибки для отладки
        print(f"Ошибка при обновлении пользователя: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Ошибка сервера: {str(e)}")

# Application endpoints
@app.post("/applications/", response_model=schemas.Application)
async def create_application(
    application: schemas.ApplicationCreate,
    db: Session = Depends(get_db)
):
    # Создаем объект заявки
    db_application = models.Application(
        **application.dict(),
        status="новый"
    )
    db.add(db_application)
    db.commit()
    db.refresh(db_application)
    
    # Создаем уведомления для всех администраторов
    admins = db.query(models.User).filter(models.User.role == "admin").all()
    for admin in admins:
        notification = models.Notification(
            user_id=admin.id,
            message=f"Новая заявка от {db_application.full_name}",
            is_read=False
        )
        db.add(notification)
    db.commit()
    
    # Отправляем уведомление всем подключенным клиентам
    await manager.broadcast({
        "type": "new_application",
        "application": {
            "id": db_application.id,
            "full_name": db_application.full_name,
            "birth_date": db_application.birth_date,
            "position": db_application.position,
            "specialization": db_application.specialization,
            "education": db_application.education,
            "citizenship": db_application.citizenship,
            "experience": db_application.experience,
            "city": db_application.city,
            "phone": db_application.phone,
            "telegram": db_application.telegram,
            "languages": db_application.languages,
            "source": db_application.source,
            "created_at": db_application.created_at.isoformat(),
            "resume_file_path": db_application.resume_file_path
        }
    })
    
    return db_application

@app.get("/applications/", response_model=List[schemas.Application])
def read_applications(
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    applications = db.query(models.Application).offset(skip).limit(limit).all()
    return applications

@app.get("/applications/{application_id}", response_model=schemas.Application)
def read_application(
    application_id: int,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    application = db.query(models.Application).filter(models.Application.id == application_id).first()
    if application is None:
        raise HTTPException(status_code=404, detail="Application not found")
    
    # Добавляем логирование
    print(f"Application data: {application.__dict__}")
    print(f"Resume file path: {application.resume_file_path}")
    
    return application

@app.put("/applications/{application_id}/status", response_model=schemas.Application)
async def update_application_status(
    application_id: int,
    status_update: schemas.StatusUpdate,
    current_user: schemas.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    application = db.query(models.Application).filter(models.Application.id == application_id).first()
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")

    # Update application status
    application.status = status_update.status

    # Create status history entry
    status_history = models.StatusHistory(
        application_id=application_id,
        status=status_update.status,
        comment=status_update.comment,
        created_by=current_user.full_name
    )
    db.add(status_history)
    db.commit()
    db.refresh(application)
    
    return application

@app.get("/applications/{application_id}/history", response_model=List[schemas.StatusHistory])
async def get_application_history(
    application_id: int,
    current_user: schemas.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    history = db.query(models.StatusHistory)\
        .filter(models.StatusHistory.application_id == application_id)\
        .order_by(models.StatusHistory.created_at.desc())\
        .all()
    return history 

@app.post("/notify")
async def notify_clients(message: dict):
    """
    Отправляет уведомления всем подключенным клиентам через WebSocket.
    Поддерживает различные типы уведомлений: new_message, new_task, new_vacancy, vacancy_assignment, etc.
    """
    await manager.broadcast(message)
    return {"status": "success"}

@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    token = websocket.query_params.get("token")
    
    if not token:
        # В случае отсутствия токена, закрываем соединение без вызова accept()
        await websocket.close(code=1008, reason="Unauthorized: No token provided")
        return
    
    try:
        # Проверяем токен
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email = payload.get("sub")
        if email is None:
            await websocket.close(code=1008, reason="Invalid token")
            return
        
        # Если токен валидный, принимаем соединение
        await websocket.accept()
        await manager.connect(websocket, email)
        
        # Отправляем подтверждение успешного подключения
        await websocket.send_json({"status": "connected", "user": email})
        
        try:
            # Ждем сообщений от клиента
            while True:
                data = await websocket.receive_text()
                try:
                    # Пробуем обработать полученные данные как JSON
                    message_data = json.loads(data)
                    # В будущем здесь можно обрабатывать команды от клиента
                    if message_data.get("type") == "ping":
                        await websocket.send_json({"type": "pong"})
                except json.JSONDecodeError:
                    # Если это не JSON, игнорируем
                    pass
        except WebSocketDisconnect:
            manager.disconnect(websocket)
    except JWTError as e:
        print(f"JWT Error in WebSocket: {e}")
        await websocket.close(code=1008, reason="Invalid or expired token")
    except Exception as e:
        print(f"WebSocket error: {e}")
        try:
            await websocket.close(code=1011, reason=f"Error: {str(e)}")
        except:
            # Если соединение уже закрыто, игнорируем ошибку
            pass

# Chat endpoints
@app.post("/applications/{application_id}/messages/", response_model=schemas.Message)
async def create_message(
    application_id: int,
    message: schemas.MessageCreate,
    current_user: schemas.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    application = db.query(models.Application).filter(models.Application.id == application_id).first()
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
    
    db_message = models.Message(
        **message.dict(),
        application_id=application_id,
        sender='hr'
    )
    db.add(db_message)
    db.commit()
    db.refresh(db_message)
    
    # Отправляем сообщение кандидату через бота
    success = await send_message_to_candidate(application_id, message.content)
    if not success:
        # Если не удалось отправить сообщение, добавляем пометку в базу
        db_message.delivery_status = 'failed'
        db.commit()
    
    return db_message

@app.get("/applications/{application_id}/messages/", response_model=List[schemas.Message])
def read_messages(
    application_id: int,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    messages = db.query(models.Message).filter(
        models.Message.application_id == application_id
    ).order_by(models.Message.created_at.asc()).all()
    return messages

# Mark messages as read
@app.post("/applications/{application_id}/messages/read")
def mark_messages_as_read(
    application_id: int,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    db.query(models.Message).filter(
        models.Message.application_id == application_id,
        models.Message.is_read == False
    ).update({"is_read": True})
    db.commit()
    return {"status": "success"} 

# Инициализация парсера HeadHunter
hh_parser = HeadHunterParser()

@app.post("/hh/search")
async def search_resumes(
    query: str,
    experience: Optional[str] = None,
    salary: Optional[int] = None,
    page: int = 0,
    current_user: schemas.User = Depends(get_current_user)
):
    try:
        resumes = await hh_parser.search_resumes(
            query=query,
            experience=experience,
            salary=salary,
            page=page
        )
        return {"resumes": resumes}
    except Exception as e:
        print(f"Ошибка при поиске резюме: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/hh/resume/{resume_id}")
async def get_resume_detail(resume_id: str):
    try:
        resume = await hh_parser.get_resume_detail(resume_id)
        if not resume:
            raise HTTPException(status_code=404, detail="Резюме не найдено")
        return resume
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/applications/{application_id}")
async def delete_application(
    application_id: int,
    current_user: schemas.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    try:
        # Находим заявку
        application = db.query(models.Application).filter(models.Application.id == application_id).first()
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")
        
        # Если есть файл резюме, удаляем его
        if application.resume_file_path:
            resume_path = RESUME_DIR / application.resume_file_path
            if resume_path.exists():
                resume_path.unlink()
        
        # Удаляем все связанные сообщения
        db.query(models.Message).filter(models.Message.application_id == application_id).delete()
        
        # Удаляем всю историю статусов
        db.query(models.StatusHistory).filter(models.StatusHistory.application_id == application_id).delete()
        
        # Удаляем саму заявку
        db.query(models.Application).filter(models.Application.id == application_id).delete()
        
        # Фиксируем изменения
        db.commit()
        
        return {"status": "success", "message": "Application deleted successfully"}
    except Exception as e:
        db.rollback()
        logger.error(f"Error deleting application: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# Улучшенный эндпоинт для получения резюме
@app.get("/resumes/{file_path:path}")
async def get_resume(file_path: str):
    # Нормализуем путь и проверяем на безопасность
    try:
        resume_path = RESUME_DIR / file_path
        resume_path = resume_path.resolve()
        if not str(resume_path).startswith(str(RESUME_DIR.resolve())):
            raise HTTPException(status_code=400, detail="Invalid file path")
        
        if not resume_path.exists():
            raise HTTPException(status_code=404, detail="Resume not found")
        
        # Определяем тип файла
        content_type = 'application/pdf'
        if resume_path.suffix.lower() == '.doc':
            content_type = 'application/msword'
        elif resume_path.suffix.lower() == '.docx':
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
        headers = {
            'Content-Disposition': f'inline; filename=\"{resume_path.name}\"',
            'Accept-Ranges': 'bytes',
            'X-Frame-Options': 'SAMEORIGIN',
            'X-Content-Type-Options': 'nosniff',
            'Content-Security-Policy': "frame-ancestors 'self';"
        }
            
        return FileResponse(
            str(resume_path),
            media_type=content_type,
            headers=headers
        )
    except Exception as e:
        logging.error(f"Error serving resume file: {str(e)}")
        raise HTTPException(status_code=500, detail="Error serving resume file")

# Монтируем директорию с резюме
app.mount("/resumes", StaticFiles(directory=str(RESUME_DIR)), name="resumes")

def extract_text_from_pdf(file_path):
    """Извлекает текст из PDF файла"""
    text = ""
    try:
        print("Начинаем извлечение текста из PDF")
        with open(file_path, 'rb') as file:
            pdf = PyPDF2.PdfReader(file)
            print(f"Количество страниц в PDF: {len(pdf.pages)}")
            for i, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                print(f"Страница {i}: извлечено {len(page_text)} символов")
                text += page_text + "\n"
        return text
    except Exception as e:
        print(f"Ошибка при извлечении текста из PDF: {str(e)}")
        raise

def extract_text_from_doc(file_path):
    """Извлекает текст из DOC/DOCX файла"""
    text = ""
    try:
        print("Начинаем извлечение текста из Word документа")
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        print(f"Извлечено {len(text)} символов из Word документа")
        return text
    except Exception as e:
        print(f"Ошибка при извлечении текста из Word: {str(e)}")
        raise
        
@app.post("/hh/parse-resume")
async def parse_resume(
    resume: UploadFile = File(...),
    current_user: schemas.User = Depends(get_current_user)
):
    """Парсит загруженное резюме и возвращает извлеченную информацию"""
    try:
        start_time = time.time()
        print("\n=== Начало обработки файла ===")
        print(f"Имя файла: {resume.filename}")
        print(f"Тип файла: {resume.content_type}")
        
        # Сохраняем файл во временный файл
        temp_file_path = ""
        try:
            # Генерируем временный файл
            temp_file_handle, temp_file_path = tempfile.mkstemp(suffix='.' + resume.filename.split('.')[-1])
            os.close(temp_file_handle)
            print(f"Создан временный файл: {temp_file_path}")
            
            # Сохраняем содержимое загруженного файла во временный файл
            with open(temp_file_path, "wb") as temp_file:
                temp_file.write(await resume.read())
            print("Файл успешно сохранен во временный файл")
                
            # Создаем директорию для сохранения резюме, если она не существует
            if not os.path.exists("resumes"):
                os.makedirs("resumes")
                
            # Сохраняем резюме в постоянное хранилище с текущей датой
            file_type = resume.filename.split('.')[-1]
            # Очищаем имя файла от спецсимволов
            clean_filename = os.path.basename(resume.filename)
            
            # Сохраняем файл в папку resumes
            current_datetime = datetime.now().strftime("%Y%m%d_%H%M%S")
            save_path = f"resumes\\{current_datetime}_{clean_filename}"
            shutil.copy(temp_file_path, save_path)
            print(f"Резюме сохранено в: {save_path}")
            
            # Парсим резюме и извлекаем текст
            print("\n=== Извлечение текста из файла ===")
            text = ""
            try:
                # Выбираем метод извлечения текста в зависимости от типа файла
                if file_type.lower() in ["pdf"]:
                    text = extract_text_from_pdf(temp_file_path)
                elif file_type.lower() in ["doc", "docx"]:
                    text = extract_text_from_doc(temp_file_path)
                else:
                    print(f"Неподдерживаемый тип файла: {file_type}")
                    return JSONResponse(
                        status_code=400,
                        content={
                            "detail": f"Неподдерживаемый тип файла: {file_type}. Поддерживаются только PDF, DOC и DOCX."
                        }
                    )
            except Exception as e:
                print(f"Ошибка при извлечении текста: {str(e)}")
                return JSONResponse(
                    status_code=400,
                    content={"detail": f"Ошибка при извлечении текста: {str(e)}"}
                )
            
            # Печатаем первые 500 символов текста для отладки
            print("\nПервые 500 символов текста:")
            print("=" * 50)
            print(text[:500])
            print("=" * 50)
            
            # Разбиваем текст на строки и убираем пустые
            lines = text.split('\n')
            non_empty_lines = [line.strip() for line in lines if line.strip()]
            print(f"\nКоличество непустых строк: {len(non_empty_lines)}")
            
            # Поиск ФИО
            print("\n=== Поиск ФИО ===")
            names = []
            print("Проверяем первые 15 строк:")
            
            # Сначала проверяем первую строку документа - часто содержит ФИО без маркеров
            if non_empty_lines and len(non_empty_lines) > 0:
                first_line = non_empty_lines[0].strip()
                print(f"Проверка первой строки на ФИО: {first_line}")
                
                # Шаблон для имени: "Фамилия Имя Отчество"
                name_pattern = r'^[А-Я][а-я]+\s+[А-Я][а-я]+(?:\s+[А-Я][а-я]+)?$'
                if re.match(name_pattern, first_line):
                    print(f"Найдено ФИО в первой строке: {first_line}")
                    names.append((first_line, 0.95))  # Высокая вероятность, если имя в первой строке
            
            # Затем ищем с использованием маркеров
            for i, line in enumerate(non_empty_lines[:15], 1):
                print(f"Строка {i}: {line}")
                
                # Проверяем наличие маркеров
                for marker in ['резюме', 'анкета', 'фио']:
                    if marker in line.lower():
                        print(f"Найден маркер '{marker}' в строке: {line}")
                        cleaned_name = clean_full_name(line)
                        if cleaned_name:
                            names.append((cleaned_name, 0.9))
            
            # Выбираем наиболее вероятное имя
            if names:
                most_probable_name = max(names, key=lambda x: x[1])
                print(f"Найдено наиболее вероятное ФИО: {most_probable_name[0]} (вероятность: {most_probable_name[1]})")
            else:
                print("ФИО не найдено в документе")
            
            # Поиск должности
            print("\n=== Поиск должности ===")
            positions = []
            
            for i, line in enumerate(non_empty_lines):
                line_lower = line.lower()
                
                if 'желаемая должность' in line_lower:
                    print(f"\nНайден маркер 'желаемая должность' в строке {i+1}:")
                    print(f"Полная строка: {line}")
                    # Извлекаем текст после маркера
                    position_text = None
                    
                    # Проверяем, не содержит ли уже текущая строка должность
                    marker_position = line_lower.find('желаемая должность')
                    remaining_text = line[marker_position + len('желаемая должность'):].strip()
                    unwanted_phrases = ['и зарплата', 'зарплата', 'занятость', 'график']
                    
                    if remaining_text and len(remaining_text) > 2 and not any(phrase in remaining_text.lower() for phrase in unwanted_phrases):
                        position_text = remaining_text
                        print(f"Извлеченный текст после маркера 'желаемая должность': '{position_text}'")
                    
                    # Если должность не найдена в текущей строке, проверяем следующую строку
                    if not position_text and i + 1 < len(non_empty_lines):
                        next_line = non_empty_lines[i + 1].strip()
                        # Проверяем, что следующая строка не содержит нежелательные фразы
                        if next_line and len(next_line) > 2 and not any(phrase in next_line.lower() for phrase in unwanted_phrases):
                            print(f"Следующая строка (содержит должность): '{next_line}'")
                            position_text = next_line
                    
                    if position_text:
                        positions.append((position_text, 0.9))
                
                elif 'должность' in line_lower and 'желаемая' not in line_lower:
                    print(f"\nНайден маркер 'должность' в строке {i+1}:")
                    print(f"Полная строка: {line}")
                    # Извлекаем текст после маркера
                    marker_position = line_lower.find('должность')
                    remaining_text = line[marker_position + len('должность'):].strip()
                    # Проверяем, что оставшийся текст не содержит нежелательные фразы
                    unwanted_phrases = ['и зарплата', 'зарплата', 'занятость', 'график']
                    if remaining_text and len(remaining_text) > 2 and not any(phrase in remaining_text.lower() for phrase in unwanted_phrases):
                        print(f"Извлеченный текст после маркера: '{remaining_text}'")
                        positions.append((remaining_text, 0.7))
                    elif i + 1 < len(non_empty_lines):
                        next_line = non_empty_lines[i + 1].strip()
                        # Проверяем, что следующая строка не содержит нежелательные фразы
                        if next_line and len(next_line) > 2 and not any(phrase in next_line.lower() for phrase in unwanted_phrases):
                            print(f"Следующая строка (содержит должность): '{next_line}'")
                            positions.append((next_line, 0.9))
            
            if positions:
                most_probable_position = max(positions, key=lambda x: x[1])
                print(f"Итоговая найденная должность: {most_probable_position[0]} (вероятность: {most_probable_position[1]})")
            else:
                print("Должность не найдена")
            
            # Поиск даты рождения
            print("\n=== Поиск даты рождения ===")
            birth_date = extract_birth_date(text)
            
            # Поиск телефона
            print("\n=== Поиск телефона ===")
            phone = None
            
            # Проверяем первые 20 строк на наличие телефона
            for i, line in enumerate(non_empty_lines[:20]):
                if any(marker in line.lower() for marker in ['тел', 'моб', 'phone', '+', '998']):
                    print(f"Найдена строка с возможным телефоном ({i+1}): {line}")
                    # Очищаем телефон
                    cleaned_phone = clean_phone(line)
                    if cleaned_phone:
                        phone = cleaned_phone
                        print(f"Извлечен телефон: {phone}")
                        break
            
            # Если телефон не найден, ищем числа похожие на телефон в первых 20 строках
            if not phone:
                for i, line in enumerate(non_empty_lines[:20]):
                    if re.search(r'\d{3,}', line):
                        print(f"Найдена строка с цифрами ({i+1}): {line}")
                        cleaned_phone = clean_phone(line)
                        if cleaned_phone:
                            phone = cleaned_phone
                            print(f"Извлечен телефон: {phone}")
                            break
            
            # Поиск email
            print("\n=== Поиск email ===")
            email = None
            # Стандартный шаблон для email
            simple_email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
            # Шаблон для email с пробелами
            spaced_email_pattern = r'[a-zA-Z0-9._%+-]+\s*@\s*[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
            
            # Сначала ищем в первых 20 строках
            for i, line in enumerate(non_empty_lines[:20]):
                # Проверяем стандартный шаблон
                email_matches = re.findall(simple_email_pattern, line)
                if email_matches:
                    email = email_matches[0]
                    print(f"Найден email (стандартный шаблон): {email}")
                    break
                
                # Проверяем шаблон с пробелами
                spaced_matches = re.findall(spaced_email_pattern, line)
                if spaced_matches:
                    # Удаляем пробелы из найденного email
                    email = re.sub(r'\s+', '', spaced_matches[0])
                    print(f"Найден email с пробелами: {spaced_matches[0]}, очищенный: {email}")
                    break
            
            # Если email не найден, ищем строки с маркерами email
            if not email:
                email_markers = ['email', 'e-mail', 'почта', 'электронная почта', 'электронный адрес']
                for i, line in enumerate(non_empty_lines[:20]):
                    line_lower = line.lower()
                    if any(marker in line_lower for marker in email_markers):
                        print(f"Найдена строка с маркером email ({i+1}): {line}")
                        
                        # Ищем стандартный email в этой строке
                        email_matches = re.findall(simple_email_pattern, line)
                        if email_matches:
                            email = email_matches[0]
                            print(f"Найден email после маркера: {email}")
                            break
                        
                        # Ищем email с пробелами
                        spaced_matches = re.findall(spaced_email_pattern, line)
                        if spaced_matches:
                            email = re.sub(r'\s+', '', spaced_matches[0])
                            print(f"Найден email с пробелами после маркера: {email}")
                            break
                        
                        # Проверяем следующую строку, если текущая строка только содержит маркер
                        if i + 1 < len(non_empty_lines):
                            next_line = non_empty_lines[i + 1]
                            email_matches = re.findall(simple_email_pattern, next_line)
                            if email_matches:
                                email = email_matches[0]
                                print(f"Найден email в следующей строке: {email}")
                                break
                            
                            spaced_matches = re.findall(spaced_email_pattern, next_line)
                            if spaced_matches:
                                email = re.sub(r'\s+', '', spaced_matches[0])
                                print(f"Найден email с пробелами в следующей строке: {email}")
                                break
            
            # Поиск образования
            print("\n=== Поиск образования ===")
            education = extract_education(text)
            print("Найдена информация об образовании:")
            for entry in education:
                for key, value in entry.items():
                    print(f"  {key}: {value}")
            
            # Поиск опыта работы
            print("\n=== Поиск опыта работы ===")
            experience = extract_experience(text)
            print("Найдена информация об опыте работы:")
            for entry in experience:
                for key, value in entry.items():
                    print(f"  {key}: {value}")
            
            # Поиск навыков
            print("\n=== Поиск навыков ===")
            skills = extract_skills(text)
            print("Найдена информация о навыках:")
            for skill in skills:
                print(f"  {skill}")
            
            # Поиск языков
            print("\n=== Поиск языков ===")
            languages = extract_languages(text)
            print("Найдена информация о языках:")
            for lang in languages:
                print(f"  {lang['language']} — {lang['level']}")
            
            # Формируем и возвращаем результат
            result = {
                "full_name": names[0][0] if names else "Не распознано",
                "birth_date": birth_date or "Не указана",
                "phone": phone or "Не указан",
                "position": positions[0][0] if positions else "Не распознана",
                "resume_file_path": save_path,
                # Дополнительные данные для отладки
                "names": names,
                "positions": positions,
                "education": education,
                "experience": experience,
                "skills": skills,
                "languages": languages,
            }
            
            # Добавляем email, если найден
            if email:
                result["email"] = email
                print(f"Email: {email}")
            
            # Добавляем метаданные для отладки
            result["_metadata"] = {
                "file_type": resume.content_type,
                "original_filename": resume.filename,
                "processing_time": time.time() - start_time,
                "text_length": len(text),
                "lines_count": len(non_empty_lines),
            }
            
            # Удаляем временный файл
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
                print("Временный файл удален")
            
            return result
        
        finally:
            # Удаляем временный файл, если он существует
            if temp_file_path and os.path.exists(temp_file_path):
                os.remove(temp_file_path)
                
    except Exception as e:
        print(f"Ошибка при обработке файла: {str(e)}")
        # Если ошибка связана с datetime
        if "datetime" in str(e):
            return JSONResponse(
                status_code=400,
                content={"detail": "Ошибка при обработке даты. Пожалуйста, сообщите администратору системы."}
            )
        return JSONResponse(
            status_code=400,
            content={"detail": f"Ошибка при обработке файла: {str(e)}"}
        )

@app.get("/hh-candidates/")
async def get_hh_candidates(
    position: Optional[str] = None,
    start_date: Optional[datetime] = None, # Добавляем параметр start_date
    end_date: Optional[datetime] = None,   # Добавляем параметр end_date
    current_user: schemas.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    query = db.query(models.HHCandidate).order_by(models.HHCandidate.created_at.desc())
    
    # Фильтрация по должности
    if position:
        # Используем ILIKE для регистронезависимого поиска с частичным совпадением
        query = query.filter(models.HHCandidate.position.ilike(f"%{position}%"))
    
    # Фильтрация по дате создания
    if start_date:
        query = query.filter(models.HHCandidate.created_at >= start_date)
    if end_date:
        # Добавляем +1 день к end_date для включения всего дня
        # end_date_inclusive = end_date + timedelta(days=1)
        # query = query.filter(models.HHCandidate.created_at < end_date_inclusive)
        # Используем <= для включения последнего дня
        query = query.filter(models.HHCandidate.created_at <= end_date)
    
    candidates = query.all()
    
    for candidate in candidates:
        # Получаем историю статусов для каждого кандидата
        candidate.status_history = (
            db.query(models.HHStatusHistory)
            .filter(models.HHStatusHistory.candidate_id == candidate.id)
            .order_by(models.HHStatusHistory.created_at.desc())
            .all()
        )

    return candidates

@app.put("/hh-candidates/{candidate_id}/status", response_model=schemas.HHCandidate)
async def update_hh_candidate_status(
    candidate_id: int,
    status_update: schemas.StatusUpdate,
    current_user: schemas.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    try:
        candidate = db.query(models.HHCandidate).filter(models.HHCandidate.id == candidate_id).first()
        if not candidate:
            raise HTTPException(status_code=404, detail="Кандидат не найден")

        # Обновляем статус кандидата
        candidate.status = status_update.status

        # Создаем запись в истории статусов
        status_history = models.HHStatusHistory(
            candidate_id=candidate_id,
            status=status_update.status,
            comment=status_update.comment,
            created_by=current_user.full_name
        )
        db.add(status_history)
        db.commit()
        db.refresh(candidate)
        
        return candidate
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=str(e))

def is_name_word(word):
    # Проверяем, что слово начинается с заглавной буквы и содержит только буквы
    return (len(word) > 1 and 
            word[0].isupper() and 
            any(c.islower() for c in word[1:]) and
            all(c.isalpha() or c.isspace() for c in word))

def name_probability(text):
    # Увеличиваем вероятность, если текст похож на полное имя
    words = text.strip().split()
    if not words:
        return 0.0
    
    # Проверяем каждое слово на соответствие формату имени
    valid_words = [w for w in words if is_name_word(w)]
    
    probability = 0.0
    if len(valid_words) >= 2:  # Имя и Фамилия минимум
        probability = 0.5
    if len(valid_words) >= 3:  # Полное имя
        probability = 0.8
    
    # Дополнительные проверки
    if any(marker in text.lower() for marker in ['резюме', 'анкета', 'фио']):
        probability += 0.1
    
    return min(1.0, probability)

def clean_full_name(text):
    if not text or not isinstance(text, str):
        return None
    
    # Список маркеров, которые нужно удалить
    markers = ['резюме обновлено', 'резюме', 'анкета', 'фио', '•']
    
    # Преобразуем текст в нижний регистр для поиска маркеров
    lower_text = text.lower()
    
    # Очищаем от маркеров
    for marker in markers:
        if marker in lower_text:
            # Найдем индекс маркера в исходном тексте (с учетом регистра)
            marker_index = lower_text.find(marker)
            marker_length = len(marker)
            
            # Получаем текст после маркера
            text = text[marker_index + marker_length:].strip()
            lower_text = text.lower()  # Обновляем и нижний регистр
    
    # Удаление дат и времени (например, "13 января 2025 в 21:17")
    text = re.sub(r'\d{1,2}\s+[а-я]+\s+\d{4}(?:\s+в\s+\d{1,2}:\d{2})?', '', text, flags=re.IGNORECASE).strip()
    
    # Ищем в тексте шаблон имени "Имя Отчество Фамилия" или "Фамилия Имя Отчество"
    name_pattern = r'([А-Я][а-я]+(?:-[А-Я][а-я]+)?\s+[А-Я][а-я]+(?:-[А-Я][а-я]+)?\s+[А-Я][а-я]+(?:-[А-Я][а-я]+)?)'
    name_match = re.search(name_pattern, text)
    
    if name_match:
        return name_match.group(1)
    
    # Если имя не найдено по шаблону, проверяем, содержит ли текст только имя (без лишних символов)
    # Удаляем все цифры, знаки препинания и лишние пробелы
    cleaned_text = re.sub(r'[0-9.,;:!?()"\']', '', text).strip()
    words = cleaned_text.split()
    
    # Проверяем, что все слова начинаются с заглавной буквы (как имена)
    if 2 <= len(words) <= 3 and all(word[0].isupper() for word in words if word):
        return ' '.join(words)
    
    return None

def extract_birth_date(text):
    """Извлекает дату рождения из текста резюме"""
    if not text:
        return None
    
    print("Начинаем поиск даты рождения...")
    
    # Преобразуем текст в список строк
    lines = text.split('\n')
    candidate_date = None
    
    # Список ключевых слов для поиска даты рождения
    birth_keywords = ['родился', 'родилась', 'дата рождения', 'год рождения', 'день рождения']
    month_names = {
        'января': '01', 'февраля': '02', 'марта': '03', 'апреля': '04',
        'мая': '05', 'июня': '06', 'июля': '07', 'августа': '08',
        'сентября': '09', 'октября': '10', 'ноября': '11', 'декабря': '12',
        'янв': '01', 'фев': '02', 'мар': '03', 'апр': '04',
        'май': '05', 'июн': '06', 'июл': '07', 'авг': '08',
        'сен': '09', 'окт': '10', 'ноя': '11', 'дек': '12'
    }
    
    # Шаг 1: Сначала ищем строки, содержащие ключевые слова
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
            
        # Проверяем наличие ключевых слов
        line_lower = line.lower()
        if any(keyword in line_lower for keyword in birth_keywords):
            print(f"Найдена строка с возможной датой рождения (строка {i+1}): {line}")
            
            # Расширенные паттерны для различных форматов дат
            patterns = [
                # "родился 3 марта 1998" или "родилась 29 сентября 1892"
                r'род(?:ился|илась)[^\d]*(\d{1,2})\s+([а-яА-Я]+)\s+(\d{4})',
                
                # "родился 03.03.1998"
                r'род(?:ился|илась)[^\d]*(\d{1,2})[\.\/](\d{1,2})[\.\/](\d{4})',
                
                # Извлечение возраста и даты из формата "Мужчина, 30 лет, родился 1 января 1990"
                r'(?:.*?,\s*\d+\s*(?:год|лет|года)[^\d]*)[^\d]*род(?:ился|илась)[^\d]*(\d{1,2})\s+([а-яА-Я]+)\s+(\d{4})',
                
                # "дата рождения: 03.03.1998"
                r'(?:дата|день|год)\s+рождения[\s:]+(\d{1,2})[\.\/](\d{1,2})[\.\/](\d{4})',
                
                # "дата рождения: 3 марта 1998"
                r'(?:дата|день|год)\s+рождения[\s:]+(\d{1,2})\s+([а-яА-Я]+)\s+(\d{4})',
                
                # Общий формат "DD.MM.YYYY" или "DD/MM/YYYY" после ключевых слов
                r'[^\d]*(\d{1,2})[\.\/](\d{1,2})[\.\/](\d{4})',
                
                # Формат "DD месяц YYYY" после ключевых слов
                r'[^\d]*(\d{1,2})\s+([а-яА-Я]+)\s+(\d{4})'
            ]
            
            for pattern in patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    print(f"Паттерн '{pattern}' сработал!")
                    if len(match.groups()) == 3:
                        day, month, year = match.groups()
                        print(f"Извлечено: день={day}, месяц={month}, год={year}")
                        
                        # Если месяц - слово, преобразуем его в число
                        if not month.isdigit():
                            month_lower = month.lower()
                            if month_lower in month_names:
                                month = month_names[month_lower]
                                print(f"Преобразовали месяц '{month_lower}' в '{month}'")
                            else:
                                print(f"Неизвестный месяц: '{month}'")
                                # Если месяц не распознан, продолжаем поиск
                                continue
                        
                        # Нормализуем день и месяц (добавляем ведущий ноль)
                        day = day.zfill(2)
                        month = month.zfill(2)
                        
                        # Проверяем валидность даты
                        try:
                            # Создаем объект datetime для проверки
                            date_obj = datetime(int(year), int(month), int(day))
                            
                            # Проверяем, что год рождения реалистичен (между 1900 и текущим годом)
                            current_year = datetime.now().year
                            if 1900 <= int(year) <= current_year:
                                print(f"Найдена дата рождения: {day}.{month}.{year}")
                                return f"{day}.{month}.{year}"
                        except ValueError:
                            # Невалидная дата, продолжаем поиск
                            continue
    
    # Шаг 2: Если не нашли по ключевым словам, ищем даты во всех строках
    print("Поиск дат рождения во всех строках...")
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
            
        # Те же паттерны, что и выше, но без привязки к ключевым словам о рождении
        date_patterns = [
            r'(\d{1,2})[\.\/](\d{1,2})[\.\/](\d{4})',
            r'(\d{1,2})\s+([а-яА-Я]+)\s+(\d{4})'
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                print(f"Найдена потенциальная дата в строке {i+1}: {line}")
                if len(match.groups()) == 3:
                    day, month, year = match.groups()
                    print(f"Извлечено: день={day}, месяц={month}, год={year}")
                    
                    # Преобразование месяца-слова в число, если необходимо
                    if not month.isdigit():
                        month_lower = month.lower()
                        if month_lower in month_names:
                            month = month_names[month_lower]
                        else:
                            continue
                    
                    day = day.zfill(2)
                    month = month.zfill(2)
                    
                    try:
                        date_obj = datetime(int(year), int(month), int(day))
                        current_year = datetime.now().year
                        
                        # Для дат без контекста быть строже - очень вероятно, что это не дата рождения
                        if 1940 <= int(year) <= current_year - 18:  # Предполагаем, что кандидату не меньше 18 лет
                            print(f"Найдена потенциальная дата рождения: {day}.{month}.{year}")
                            candidate_date = f"{day}.{month}.{year}"
                    except ValueError:
                        continue
    
    return candidate_date

def clean_phone(text):
    """Очищает и форматирует телефонный номер"""
    if not text:
        return None
    
    # Проверяем, что строка не содержит нежелательные тексты
    unwanted_texts = ['желательное время', 'не имеет значения', 'занятость', 'опыт работы']
    for unwanted in unwanted_texts:
        if unwanted.lower() in text.lower():
            return None
    
    # Шаблоны для поиска номеров телефонов
    phone_patterns = [
        r'\+?\d{1,3}\s?\(?\d{2,3}\)?\s?\d{3}[\s-]?\d{2}[\s-]?\d{2}',  # +998 (99) 999-99-99 или 998 99 999 99 99
        r'\d{2,3}[\s-]?\d{3}[\s-]?\d{2}[\s-]?\d{2}',  # 99 999 99 99
        r'\+?\d{10,12}'  # +998999999999 или 998999999999
    ]
    
    # Ищем телефоны по шаблонам
    for pattern in phone_patterns:
        matches = re.findall(pattern, text)
        if matches:
            # Берем первое совпадение
            phone = matches[0]
            
            # Удаляем все нецифровые символы, кроме '+'
            cleaned = ''.join(c for c in phone if c.isdigit() or c == '+')
            
            # Проверяем длину номера после очистки
            if len(cleaned) >= 9:  # Минимальная длина для номера
                # Форматируем номер в стандартный вид
                if cleaned.startswith('+'):
                    # Уже есть префикс
                    pass
                elif len(cleaned) >= 12:
                    # Добавляем префикс '+' к длинному номеру
                    cleaned = '+' + cleaned
                elif len(cleaned) >= 9 and len(cleaned) <= 11:
                    # Это вероятно номер без кода страны, добавляем +998
                    cleaned = '+998' + cleaned[-9:]
                
                # Проверяем, что номер начинается с +998 (Узбекистан) или другого валидного префикса
                # Если номер не начинается с +, добавим +
                if not cleaned.startswith('+'):
                    cleaned = '+' + cleaned
                
                # Форматируем номер для удобочитаемости (если это номер Узбекистана)
                if cleaned.startswith('+998') and len(cleaned) >= 12:
                    formatted = f"+998 ({cleaned[4:6]}) {cleaned[6:9]}-{cleaned[9:11]}-{cleaned[11:13]}"
                    return formatted
                
                return cleaned
    
    return None

def extract_education(text):
    """Извлекает информацию об образовании"""
    education_data = []
    education_markers = ['образование', 'учебное заведение', 'университет', 'институт', 'колледж']
    
    current_entry = {}
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        # Ищем год
        year_match = re.search(r'\b(19|20)\d{2}\b', line)
        
        # Ищем учебное заведение
        is_education_line = any(marker.lower() in line.lower() for marker in education_markers)
        
        if year_match or is_education_line:
            if current_entry and (year_match or is_education_line):
                education_data.append(current_entry.copy())
                current_entry = {}
            
            if year_match:
                current_entry['year'] = year_match.group(0)
            
            # Очищаем строку от года для извлечения названия
            institution = re.sub(r'\b(19|20)\d{2}\b', '', line).strip()
            if institution and is_education_line:
                current_entry['institution'] = institution
    
    if current_entry:
        education_data.append(current_entry)
    
    return education_data

def extract_experience(text):
    experience_data = []
    experience_started = False
    current_entry = {}
    
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if 'опыт работы' in line.lower():
            experience_started = True
            # Извлекаем общий стаж
            match = re.search(r'(\d+)\s*(?:год|лет|года).*?(\d+)\s*(?:месяц|месяца|месяцев)?', line)
            if match:
                years, months = match.groups()
                current_entry['total_experience'] = f"{years} лет {months} месяцев"
            continue
            
        if experience_started:
            if 'образование' in line.lower() or 'навыки' in line.lower():
                break
                
            # Ищем даты
            date_match = re.search(r'([А-Яа-я]+\s+\d{4})\s*[—–-]\s*([А-Яа-я]+\s+\d{4}|настоящее время)', line)
            if date_match:
                if current_entry:
                    experience_data.append(current_entry)
                current_entry = {'period': f"{date_match.group(1)} - {date_match.group(2)}"}
                
            # Ищем компанию
            if 'ооо' in line.lower() or 'компания' in line.lower() or 'фирма' in line.lower():
                current_entry['company'] = line
                
            # Ищем должность
            if any(pos in line.lower() for pos in ['менеджер', 'специалист', 'руководитель', 'директор']):
                current_entry['position'] = line
                
    if current_entry:
        experience_data.append(current_entry)
        
    return experience_data

def extract_skills(text):
    # Implementation of extract_skills function
    # This function should return a list of skills extracted from the text
    # For now, we'll return an empty list
    return []

def extract_languages(text):
    """Извлекает информацию о языках"""
    languages = []
    language_markers = ['язык', 'languages', 'владение языками']
    language_levels = {
        'a1': 'начальный',
        'a2': 'элементарный',
        'b1': 'средний',
        'b2': 'выше среднего',
        'c1': 'продвинутый',
        'c2': 'в совершенстве',
        'native': 'родной',
        'свободно': 'в совершенстве'
    }
    
    in_language_section = False
    seen_languages = set()
    
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        # Определяем начало секции языков
        if any(marker.lower() in line.lower() for marker in language_markers):
            in_language_section = True
            continue
            
        if in_language_section:
            # Проверяем, не начался ли новый раздел
            if any(marker in line.lower() for marker in ['опыт работы', 'образование', 'навыки']):
                break
                
            # Извлекаем язык и уровень
            parts = line.split('—')
            if len(parts) >= 2:
                language = parts[0].strip()
                level = parts[-1].strip()
                
                # Проверяем, не видели ли мы уже этот язык
                if language.lower() not in seen_languages:
                    seen_languages.add(language.lower())
                    languages.append({
                        'language': language,
                        'level': level
                    })
    
    return languages

def normalize_phone(phone: str) -> str:
    """Нормализует телефонный номер, оставляя только цифры"""
    return ''.join(filter(str.isdigit, phone))

def normalize_full_name(full_name: str) -> str:
    """Нормализует полное имя (убирает лишние пробелы)"""
    return ' '.join(full_name.split())

@app.get("/api/hh/candidates/check")
async def check_hh_candidate(full_name: str, phone: str, db: Session = Depends(get_db)):
    """Проверяет, существует ли кандидат с такими же full_name и phone"""
    try:
        # Нормализуем входные данные
        clean_phone = normalize_phone(phone)
        normalized_name = normalize_full_name(full_name)
        
        logging.info(f"Проверка существования кандидата: {normalized_name} ({clean_phone})")
        
        # Ищем кандидата в базе
        candidate = db.query(models.HHCandidate).filter(
            models.HHCandidate.full_name == normalized_name,
            models.HHCandidate.phone.like(f"%{clean_phone}%")
        ).first()
        
        if candidate:
            logging.info(f"Найден существующий кандидат: ID={candidate.id}")
            return {
                "exists": True,
                "candidate": {
                    "id": candidate.id,
                    "full_name": candidate.full_name,
                    "position": candidate.position,
                    "status": candidate.status,
                    "created_at": candidate.created_at.isoformat()
                }
            }
        
        logging.info("Кандидат не найден в базе")
        return {"exists": False}
        
    except Exception as e:
        logging.error(f"Ошибка при проверке существования кандидата: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Ошибка при проверке существования кандидата"
        )

@app.post("/api/hh/candidates")
async def import_hh_candidate(
    full_name: str = Form(...),
    phone: str = Form(...),
    specialty: str = Form(...),
    birth_date: str = Form(...),
    education: str = Form(...),
    experience: str = Form(...),
    languages: str = Form(...),
    location: str = Form(...),
    citizenship: str = Form(...),
    status: Optional[str] = Form("новый"),  # Добавляем статус с дефолтным значением "новый"
    comment: Optional[str] = Form("Импортирован из HeadHunter"),  # Добавляем комментарий с дефолтным значением
    resume: UploadFile = File(None),
    current_user: Optional[schemas.User] = None,
    db: Session = Depends(get_db)
):
    """Импортирует нового кандидата из HeadHunter"""
    try:
        # Нормализуем входные данные
        normalized_name = normalize_full_name(full_name)
        clean_phone = normalize_phone(phone)
        
        # Логируем входящие данные для отладки
        logging.info(f"Попытка импорта кандидата:")
        logging.info(f"ФИО: {normalized_name}")
        logging.info(f"Телефон: {clean_phone}")
        logging.info(f"Специальность: {specialty}")
        logging.info(f"Дата рождения: {birth_date}")
        logging.info(f"Образование: {education}")
        logging.info(f"Опыт работы: {experience}")
        logging.info(f"Языки: {languages}")
        logging.info(f"Местоположение: {location}")
        logging.info(f"Гражданство: {citizenship}")
        
        # Проверяем, существует ли уже кандидат
        existing_candidate = db.query(models.HHCandidate).filter(
            models.HHCandidate.full_name == normalized_name,
            models.HHCandidate.phone.like(f"%{clean_phone}%")
        ).first()
        
        if existing_candidate:
            logging.warning(f"Попытка повторного импорта кандидата: {normalized_name} ({clean_phone})")
            return JSONResponse(
                status_code=400,
                content={
                    "success": False,
                    "error": "Кандидат уже существует в базе данных",
                    "candidate": {
                        "id": existing_candidate.id,
                        "full_name": existing_candidate.full_name,
                        "position": existing_candidate.position,
                        "status": existing_candidate.status,
                        "created_at": existing_candidate.created_at.isoformat()
                    }
                }
            )
        
        # Обрабатываем PDF-резюме, если оно есть
        resume_file_path = None
        if resume:
            # Создаем безопасное имя файла с временной меткой
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            safe_filename = f"{timestamp}_{secure_filename(resume.filename)}"
            
            # Убеждаемся, что директория для резюме существует
            os.makedirs("resumes", exist_ok=True)
            
            # Сохраняем файл
            file_path = os.path.join("resumes", safe_filename)
            with open(file_path, "wb") as f:
                f.write(await resume.read())
            
            resume_file_path = f"resumes/{safe_filename}"
            logging.info(f"Сохранено резюме: {resume_file_path}")
        
        # Создаем новую запись кандидата
        new_candidate = models.HHCandidate(
            full_name=normalized_name,
            birth_date=birth_date,
            phone=clean_phone,
            position=specialty,
            resume_file_path=resume_file_path,
            status=status
        )
        
        db.add(new_candidate)
        db.flush()  # Получаем id без коммита
        
        # Добавляем запись в историю статусов
        status_history = models.HHStatusHistory(
            candidate_id=new_candidate.id,
            status=status,
            comment=comment,
            created_by=current_user.full_name if current_user else "HeadHunter Extension"
        )
        db.add(status_history)
        
        # Сохраняем изменения
        db.commit()
        logging.info(f"Успешно импортирован кандидат с ID: {new_candidate.id}")
        
        return {
            "success": True,
            "id": new_candidate.id,
            "message": "Кандидат успешно импортирован",
            "candidate": {
                "id": new_candidate.id,
                "full_name": new_candidate.full_name,
                "position": new_candidate.position,
                "status": new_candidate.status,
                "created_at": new_candidate.created_at.isoformat()
            }
        }
        
    except Exception as e:
        logging.error(f"Ошибка при импорте кандидата: {str(e)}")
        db.rollback()
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": f"Ошибка при импорте кандидата: {str(e)}"}
        )

# Функция для извлечения упоминаний пользователей в формате @email из текста
def extract_mentions(text: str) -> Set[str]:
    """
    Извлекает упоминания пользователей в формате @email из текста
    
    Например: "Задача для @user@example.com, срочно!"
    """
    if not text:
        return set()
    
    # Регулярное выражение для поиска упоминаний в формате @email
    mention_pattern = r'@([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
    mentions = re.findall(mention_pattern, text)
    
    return set(mentions)

# Функция для создания уведомлений для упомянутых пользователей
def create_mention_notifications(db: Session, task_id: int, text: str, creator_id: int):
    """
    Создает уведомления для пользователей, упомянутых в тексте
    """
    if not text:
        return
    
    # Получаем все упоминания из текста
    mentioned_emails = extract_mentions(text)
    
    if not mentioned_emails:
        return
    
    # Получаем информацию о задаче
    task = db.query(models.Task).filter(models.Task.id == task_id).first()
    if not task:
        return
    
    # Получаем пользователей по их email
    for email in mentioned_emails:
        mentioned_user = db.query(models.User).filter(models.User.email == email).first()
        
        if mentioned_user and mentioned_user.id != creator_id:  # Не уведомляем создателя о его же упоминании
            # Проверяем, существует ли уже уведомление
            existing = db.query(models.Notification).filter(
                models.Notification.user_id == mentioned_user.id,
                models.Notification.task_id == task_id,
                models.Notification.message.like(f"%упомянул вас в задаче%")
            ).first()
            
            if not existing:
                # Создаем уведомление
                creator = db.query(models.User).filter(models.User.id == creator_id).first()
                creator_name = creator.full_name if creator else "Пользователь"
                
                notification = models.Notification(
                    user_id=mentioned_user.id,
                    task_id=task_id,
                    message=f"{creator_name} упомянул вас в задаче: {task.title}"
                )
                db.add(notification)
                print(f"Создано уведомление об упоминании для пользователя {mentioned_user.email}")
    
    db.commit()

@app.post("/tasks/", response_model=schemas.Task)
def create_task(task: schemas.TaskCreate, current_user: schemas.User = Depends(get_current_user), db: Session = Depends(get_db)):
    """
    Создание новой задачи для календаря, автоматически устанавливает текущего пользователя как создателя
    """
    from sqlalchemy.orm import joinedload
    
    # Создаем новую задачу, явно устанавливая ID текущего пользователя как создателя
    db_task = models.Task(
        title=task.title,
        description=task.description,
        date=task.date,
        time=task.time,
        status=task.status,
        assigned_to=task.assigned_to,
        created_by_id=current_user.id  # Сохраняем ID текущего пользователя как создателя
    )
    
    db.add(db_task)
    db.commit()
    db.refresh(db_task)
    
    # Проверяем наличие упоминаний в описании задачи
    combined_text = f"{task.title} {task.description}" if task.description else task.title
    create_mention_notifications(db, db_task.id, combined_text, current_user.id)
    
    # Получаем задачу с данными о создателе
    created_task = db.query(models.Task).options(joinedload(models.Task.creator)).filter(models.Task.id == db_task.id).first()
    
    # Отладочное сообщение
    print(f"Создана задача ID: {created_task.id}, Название: {created_task.title}, Создатель ID: {created_task.created_by_id}")
    if created_task.creator:
        print(f"  Данные создателя: {created_task.creator.full_name} ({created_task.creator.email})")
    
    return created_task

@app.get("/tasks/", response_model=List[schemas.Task])
def get_tasks(
    current_user: schemas.User = Depends(get_current_user),
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """
    Получение списка задач с возможной фильтрацией по диапазону дат
    """
    from sqlalchemy.orm import joinedload
    
    # Используем joinedload для eagerly loading связанных объектов
    query = db.query(models.Task).options(joinedload(models.Task.creator))
    
    if start_date:
        query = query.filter(models.Task.date >= start_date)
    
    if end_date:
        query = query.filter(models.Task.date <= end_date)
    
    # Если пользователь не администратор, то возвращаем только задачи,
    # где он является создателем, ИЛИ назначенным исполнителем, ИЛИ упомянут.
    if current_user.role != "admin":
        mention_pattern = f"%@{current_user.email}%"
        query = query.filter(
            or_(
                models.Task.assigned_to == current_user.id,
                models.Task.created_by_id == current_user.id,
                models.Task.title.ilike(mention_pattern), # Регистронезависимый поиск
                and_( # Проверяем описание только если оно не NULL
                    models.Task.description != None,
                    models.Task.description.ilike(mention_pattern)
                )
            )
        )
    
    tasks = query.order_by(models.Task.date.asc()).all()
    
    # Выводим информацию о загруженных задачах для отладки
    print(f"Загружено {len(tasks)} задач")
    for task in tasks:
        print(f"Задача ID: {task.id}, Название: {task.title}, Создатель ID: {task.created_by_id}")
        if task.creator:
            print(f"  Данные создателя: {task.creator.full_name} ({task.creator.email})")
    
    return tasks

@app.get("/tasks/{task_id}", response_model=schemas.Task)
def get_task(task_id: int, current_user: schemas.User = Depends(get_current_user), db: Session = Depends(get_db)):
    """
    Получение информации о конкретной задаче
    """
    from sqlalchemy.orm import joinedload
    
    # Используем joinedload для подгрузки связанных объектов
    db_task = db.query(models.Task).options(joinedload(models.Task.creator)).filter(models.Task.id == task_id).first()
    
    if db_task is None:
        raise HTTPException(status_code=404, detail="Задача не найдена")
    
    # Логирование перед проверкой прав
    print(f"Checking access for user ID: {current_user.id} (Role: {current_user.role}) to Task ID: {task_id}")
    print(f"  Task assigned_to: {db_task.assigned_to}")
    print(f"  Task created_by_id: {db_task.created_by_id}")
    print(f"  Task title: {db_task.title}")
    print(f"  Task description: {db_task.description}")
    
    # Проверяем, что пользователь имеет доступ к задаче
    # Администратор имеет доступ ко всем задачам.
    # Другие пользователи могут просматривать задачи, где они создатели, исполнители ИЛИ упомянуты.
    if current_user.role != "admin":
        is_assignee = db_task.assigned_to == current_user.id
        is_creator = db_task.created_by_id == current_user.id
        
        # Проверка упоминания
        mention_pattern = f"%@{current_user.email}%"
        is_mentioned_in_title = db_task.title is not None and db.query(func.count(models.Task.id)).filter(models.Task.id == task_id, models.Task.title.ilike(mention_pattern)).scalar() > 0
        is_mentioned_in_desc = db_task.description is not None and db.query(func.count(models.Task.id)).filter(models.Task.id == task_id, models.Task.description.ilike(mention_pattern)).scalar() > 0
        is_mentioned = is_mentioned_in_title or is_mentioned_in_desc

        print(f"  Is assignee? {is_assignee}")
        print(f"  Is creator? {is_creator}")
        print(f"  Is mentioned? {is_mentioned} (Title: {is_mentioned_in_title}, Desc: {is_mentioned_in_desc})")
        
        if not is_assignee and not is_creator and not is_mentioned:
            print(f"  Access denied for user ID: {current_user.id}")
            raise HTTPException(status_code=403, detail="Нет доступа к этой задаче")
        else:
            print(f"  Access granted for user ID: {current_user.id}")
    else:
        print(f"  Access granted for admin user ID: {current_user.id}")
    
    # Выводим отладочную информацию (если доступ предоставлен)
    print(f"Задача ID: {db_task.id}, Название: {db_task.title}, Создатель ID: {db_task.created_by_id}")
    
    return db_task

@app.put("/tasks/{task_id}", response_model=schemas.Task)
def update_task(
    task_id: int,
    task_update: schemas.TaskUpdate,
    current_user: schemas.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Обновление существующей задачи
    """
    from sqlalchemy.orm import joinedload
    
    # Используем joinedload для подгрузки связанных объектов
    db_task = db.query(models.Task).options(joinedload(models.Task.creator)).filter(models.Task.id == task_id).first()
    
    if db_task is None:
        raise HTTPException(status_code=404, detail="Задача не найдена")
    
    # Проверяем, что пользователь имеет доступ к задаче
    # Администратор имеет доступ ко всем задачам.
    # Другие пользователи могут обновлять только те задачи, где они создатели или исполнители.
    if current_user.role != "admin":
        if db_task.assigned_to != current_user.id and db_task.created_by_id != current_user.id:
            raise HTTPException(status_code=403, detail="Нет доступа к этой задаче")
    
    # Сохраняем оригинальный created_by_id, если новый не предоставлен
    original_creator_id = db_task.created_by_id
    
    # Обновляем поля задачи
    update_data = task_update.dict(exclude_unset=True)
    for key, value in update_data.items():
        setattr(db_task, key, value)
    
    # Убеждаемся, что created_by_id не сбросился, если не был предоставлен новый
    if 'created_by_id' not in update_data and original_creator_id:
        db_task.created_by_id = original_creator_id
    
    db.commit()
    db.refresh(db_task)
    
    # Проверяем наличие упоминаний в обновленном тексте задачи
    combined_text = f"{db_task.title} {db_task.description}" if db_task.description else db_task.title
    create_mention_notifications(db, db_task.id, combined_text, current_user.id)
    
    # Снова загружаем задачу с создателем
    db_task = db.query(models.Task).options(joinedload(models.Task.creator)).filter(models.Task.id == task_id).first()
    
    return db_task

@app.delete("/tasks/{task_id}", response_model=schemas.Task)
def delete_task(task_id: int, current_user: schemas.User = Depends(get_current_user), db: Session = Depends(get_db)):
    """
    Удаление задачи
    """
    db_task = db.query(models.Task).filter(models.Task.id == task_id).first()
    if db_task is None:
        raise HTTPException(status_code=404, detail="Задача не найдена")
    
    # Проверяем, что пользователь имеет доступ к задаче
    if current_user.role != "admin" and db_task.assigned_to != current_user.id:
        raise HTTPException(status_code=403, detail="Нет доступа к этой задаче")
    
    # Запоминаем данные задачи перед удалением
    task_data = schemas.Task.from_orm(db_task)
    
    db.delete(db_task)
    db.commit()
    
    return task_data

# ===== Уведомления =====
@app.get("/notifications/", response_model=List[schemas.Notification])
async def get_notifications(
    current_user: schemas.User = Depends(get_current_user),
    db: Session = Depends(get_db),
    include_read: bool = False
):
    """
    Получение списка уведомлений для текущего пользователя
    """
    query = db.query(models.Notification).filter(models.Notification.user_id == current_user.id)
    
    if not include_read:
        query = query.filter(models.Notification.is_read == False)
    
    notifications = query.order_by(models.Notification.created_at.desc()).all()
    return notifications

@app.post("/notifications/{notification_id}/read", response_model=dict)
async def mark_notification_as_read(
    notification_id: int,
    current_user: schemas.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Отметить уведомление как прочитанное
    """
    notification = db.query(models.Notification).filter(
        models.Notification.id == notification_id,
        models.Notification.user_id == current_user.id
    ).first()
    
    if not notification:
        raise HTTPException(status_code=404, detail="Уведомление не найдено")
    
    notification.is_read = True
    db.commit()
    
    return {"status": "success"}

@app.post("/notifications/read-all", response_model=dict)
async def mark_all_notifications_as_read(
    current_user: schemas.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Отметить все уведомления пользователя как прочитанные
    """
    db.query(models.Notification).filter(
        models.Notification.user_id == current_user.id,
        models.Notification.is_read == False
    ).update({"is_read": True})
    
    db.commit()
    
    return {"status": "success"}

# Фоновая задача для проверки и создания уведомлений
async def check_upcoming_tasks():
    """
    Проверяет задачи, которые начинаются в ближайшее время, и создает уведомления
    """
    while True:
        try:
            # Используем независимую сессию для проверки
            db = SessionLocal()
            
            try:
                # Получаем текущее время и время для проверки (+5 минут)
                now = datetime.now()
                check_time = now + timedelta(minutes=5)
                
                # Форматируем для сравнения с данными в БД
                check_date = check_time.strftime('%Y-%m-%d')
                check_time_str = check_time.strftime('%H:%M')
                
                # Округляем до минут для корректного сравнения
                check_time_rounded = check_time.replace(second=0, microsecond=0)
                now_rounded = now.replace(second=0, microsecond=0)
                
                # Ищем все запланированные задачи, которые начинаются через 5 минут
                upcoming_tasks = db.query(models.Task).filter(
                    models.Task.date == check_date,
                    models.Task.time == check_time_str,
                    models.Task.status == "planned"
                ).all()
                
                print(f"Found {len(upcoming_tasks)} tasks scheduled for {check_date} {check_time_str}")
                
                # Обрабатываем найденные задачи
                for task in upcoming_tasks:
                    # Определяем получателей уведомлений (создатель и назначенный пользователь)
                    recipients = set()
                    if task.created_by_id:
                        recipients.add(task.created_by_id)
                    if task.assigned_to:
                        recipients.add(task.assigned_to)
                    
                    # Создаем уведомление для каждого получателя
                    for user_id in recipients:
                        # Проверяем, не создавалось ли уже уведомление для этой задачи и пользователя
                        existing = db.query(models.Notification).filter(
                            models.Notification.task_id == task.id,
                            models.Notification.user_id == user_id,
                            models.Notification.created_at >= now - timedelta(minutes=10)
                        ).first()
                        
                        if not existing:
                            notification = models.Notification(
                                user_id=user_id,
                                task_id=task.id,
                                message=f"Через 5 минут начинается задача: {task.title}"
                            )
                            db.add(notification)
                            print(f"Created notification for user {user_id}, task {task.id}")
                            
                            # Отправляем WebSocket уведомление
                            # Получаем email пользователя по ID
                            user = db.query(models.User).filter(models.User.id == user_id).first()
                            if user and user.email in manager.authenticated_users:
                                print(f"Sending upcoming_task WebSocket notification to {user.email}")
                                await manager.authenticated_users[user.email].send_json({
                                    "type": "upcoming_task",
                                    "task": {
                                        "id": task.id,
                                        "title": task.title,
                                        "time": task.time,
                                        "date": task.date
                                    }
                                })
                
                # Сохраняем изменения
                db.commit()
                
            except Exception as e:
                print(f"Error checking upcoming tasks: {e}")
                db.rollback()
            finally:
                db.close()
            
            # Проверяем каждую минуту
            await asyncio.sleep(60)
            
        except Exception as e:
            print(f"Critical error in check_upcoming_tasks: {e}")
            # В случае критической ошибки ждем немного перед следующей попыткой
            await asyncio.sleep(10)

@app.on_event("startup")
async def start_background_tasks():
    """
    Запускает фоновые задачи при старте приложения
    """
    # Регистрируем обработчики для подтверждений сообщений
    try:
        from telegram_service import register_callback_handlers
        from bot import dp
        register_callback_handlers(dp)
        logging.info("Telegram confirmation callback handlers registered successfully")
    except Exception as e:
        logging.error(f"Error registering telegram confirmation handlers: {str(e)}")
    
    # Создаем и запускаем фоновую задачу для проверки предстоящих заданий
    asyncio.create_task(check_upcoming_tasks())

@app.get("/users/", response_model=List[schemas.User])
def get_users(current_user: schemas.User = Depends(get_current_user), db: Session = Depends(get_db)):
    """
    Получение списка пользователей (только для авторизованных пользователей)
    """
    users = db.query(models.User).order_by(models.User.full_name).all()
    return users

@app.get("/users/{user_id}", response_model=schemas.User)
def get_user_by_id(
    user_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """Получение информации о пользователе по ID"""
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Пользователь не найден")
    return user

# Новые API-эндпоинты для аналитики

@app.get("/analytics/status-history")
async def get_status_history_analytics(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    current_user: schemas.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Получение статистики изменений статусов заявок и кандидатов по времени.
    Возвращает данные по изменениям статусов, сгруппированные по дням.
    """
    # Преобразуем строки дат в объекты datetime, если они предоставлены
    start_datetime = None
    end_datetime = None
    
    if start_date:
        start_datetime = datetime.fromisoformat(start_date.replace('Z', '+00:00'))
    
    if end_date:
        end_datetime = datetime.fromisoformat(end_date.replace('Z', '+00:00'))
    
    # Получаем историю статусов из обеих таблиц
    application_history_query = db.query(models.StatusHistory)
    hh_history_query = db.query(models.HHStatusHistory)
    
    # Применяем фильтрацию по датам, если указаны
    if start_datetime:
        application_history_query = application_history_query.filter(
            models.StatusHistory.created_at >= start_datetime
        )
        hh_history_query = hh_history_query.filter(
            models.HHStatusHistory.created_at >= start_datetime
        )
    
    if end_datetime:
        application_history_query = application_history_query.filter(
            models.StatusHistory.created_at <= end_datetime
        )
        hh_history_query = hh_history_query.filter(
            models.HHStatusHistory.created_at <= end_datetime
        )
    
    # Получаем данные из базы
    application_history = application_history_query.all()
    hh_history = hh_history_query.all()
    
    # Структура для агрегации данных по дням и статусам
    daily_status_counts = defaultdict(lambda: defaultdict(int))
    
    # Обрабатываем историю статусов заявок
    for history_item in application_history:
        # Извлекаем дату (без времени)
        date_str = history_item.created_at.strftime('%Y-%m-%d')
        status = history_item.status.lower()
        
        # Инкрементируем счетчик для этого статуса в этот день
        daily_status_counts[date_str][status] += 1
    
    # Обрабатываем историю статусов HH кандидатов
    for history_item in hh_history:
        # Извлекаем дату (без времени)
        date_str = history_item.created_at.strftime('%Y-%m-%d')
        status = history_item.status.lower()
        
        # Инкрементируем счетчик для этого статуса в этот день
        daily_status_counts[date_str][f"hh_{status}"] += 1
    
    # Преобразуем в список для ответа API, сортируем по дате
    result = []
    for date_str, status_counts in sorted(daily_status_counts.items()):
        # Вычисляем день недели и название месяца для удобства отображения
        try:
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            weekday = calendar.day_name[date_obj.weekday()]
            month = calendar.month_name[date_obj.month]
        except:
            weekday = ""
            month = ""
        
        # Комбинированные счетчики
        combined_counts = defaultdict(int)
        for status, count in status_counts.items():
            # Если это статус HH кандидата
            if status.startswith('hh_'):
                base_status = status[3:]  # Убираем префикс 'hh_'
                combined_counts[f"combined_{base_status}"] += count
            else:
                combined_counts[f"combined_{status}"] += count
                
        # Добавляем комбинированные счетчики в результат
        status_counts.update(combined_counts)
        
        result.append({
            "date": date_str,
            "weekday": weekday,
            "month": month,
            "counts": status_counts
        })
    
    # Общая статистика
    total_stats = defaultdict(int)
    for date_data in result:
        for status, count in date_data["counts"].items():
            total_stats[status] += count
    
    return {
        "daily_data": result,
        "total_stats": dict(total_stats)
    }

# Дополнительный эндпоинт для получения помесячной статистики
@app.get("/analytics/monthly-status-history")
async def get_monthly_status_history_analytics(
    year: Optional[int] = None,
    current_user: schemas.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Получение статистики изменений статусов заявок и кандидатов по месяцам.
    """
    # Если год не указан, используем текущий
    if not year:
        year = datetime.now().year
    
    # Создаем даты начала и конца для заданного года
    start_datetime = datetime(year, 1, 1)
    end_datetime = datetime(year, 12, 31, 23, 59, 59)
    
    # Получаем историю статусов из обеих таблиц
    application_history = db.query(models.StatusHistory).filter(
        models.StatusHistory.created_at >= start_datetime,
        models.StatusHistory.created_at <= end_datetime
    ).all()
    
    hh_history = db.query(models.HHStatusHistory).filter(
        models.HHStatusHistory.created_at >= start_datetime,
        models.HHStatusHistory.created_at <= end_datetime
    ).all()
    
    # Структура для агрегации данных по месяцам и статусам
    monthly_status_counts = defaultdict(lambda: defaultdict(int))
    
    # Обрабатываем историю статусов заявок
    for history_item in application_history:
        # Извлекаем месяц
        month_str = history_item.created_at.strftime('%Y-%m')
        month_name = calendar.month_name[history_item.created_at.month]
        status = history_item.status.lower()
        
        # Инкрементируем счетчик для этого статуса в этот месяц
        monthly_status_counts[month_str]['name'] = month_name
        monthly_status_counts[month_str][status] += 1
    
    # Обрабатываем историю статусов HH кандидатов
    for history_item in hh_history:
        # Извлекаем месяц
        month_str = history_item.created_at.strftime('%Y-%m')
        month_name = calendar.month_name[history_item.created_at.month]
        status = history_item.status.lower()
        
        # Инкрементируем счетчик для этого статуса в этот месяц
        monthly_status_counts[month_str]['name'] = month_name
        monthly_status_counts[month_str][f"hh_{status}"] += 1
    
    # Преобразуем в список для ответа API, сортируем по месяцу
    result = []
    for month_str, status_counts in sorted(monthly_status_counts.items()):
        # Комбинированные счетчики
        combined_counts = defaultdict(int)
        for status, count in status_counts.items():
            if status == 'name':
                continue
                
            # Если это статус HH кандидата
            if status.startswith('hh_'):
                base_status = status[3:]  # Убираем префикс 'hh_'
                combined_counts[f"combined_{base_status}"] += count
            else:
                combined_counts[f"combined_{status}"] += count
                
        # Добавляем комбинированные счетчики в результат
        status_counts.update(combined_counts)
        
        result.append({
            "month": month_str,
            "month_name": status_counts.pop('name', ''),
            "counts": status_counts
        })
    
    return result

# Функции для работы с вакансиями

def get_vacancies(db: Session, skip: int = 0, limit: int = 100, status: Optional[str] = None):
    query = db.query(models.Vacancy)
    if status:
        query = query.filter(models.Vacancy.status == status)
    return query.order_by(models.Vacancy.created_at.desc()).offset(skip).limit(limit).all()

def get_vacancy(db: Session, vacancy_id: int):
    """
    Получить вакансию по ID с загрузкой связанных объектов.
    Используем joinedload для загрузки creator и assignments.
    """
    # Получаем вакансию с загрузкой создателя и назначений
    vacancy = db.query(models.Vacancy)\
             .options(
                joinedload(models.Vacancy.creator),
                joinedload(models.Vacancy.assignments).joinedload(models.VacancyAssignment.recruiter)
             )\
             .filter(models.Vacancy.id == vacancy_id)\
             .first()
    
    if vacancy:
        # Дополнительно загружаем рекрутеров для каждого назначения
        for assignment in vacancy.assignments:
            if assignment.recruiter_id:
                # Явно загружаем информацию о рекрутере
                recruiter = db.query(models.User).filter(models.User.id == assignment.recruiter_id).first()
                if recruiter:
                    assignment.recruiter = recruiter
    
    return vacancy

def create_vacancy(db: Session, vacancy: schemas.VacancyCreate, user_id: int):
    db_vacancy = models.Vacancy(**vacancy.model_dump(), created_by_id=user_id)
    db.add(db_vacancy)
    db.commit()
    db.refresh(db_vacancy)
    
    # Отправляем уведомления всем рекрутерам
    recruiters = db.query(models.User).filter(models.User.role == "recruiter").all()
    for recruiter in recruiters:
        notification = models.Notification(
            user_id=recruiter.id,
            message=f"Новая вакансия: {vacancy.title}",
            vacancy_id=db_vacancy.id,
            is_read=False
        )
        db.add(notification)
    
    db.commit()
    return db_vacancy

def update_vacancy(db: Session, vacancy_id: int, vacancy: schemas.VacancyUpdate):
    db_vacancy = get_vacancy(db, vacancy_id)
    if not db_vacancy:
        return None
    
    vacancy_data = vacancy.model_dump(exclude_unset=True)
    for key, value in vacancy_data.items():
        setattr(db_vacancy, key, value)
    
    db.add(db_vacancy)
    db.commit()
    db.refresh(db_vacancy)
    return db_vacancy

def delete_vacancy(db: Session, vacancy_id: int):
    db_vacancy = get_vacancy(db, vacancy_id)
    if not db_vacancy:
        return False
    
    db.delete(db_vacancy)
    db.commit()
    return True

# Функции для работы с назначениями вакансий

def get_vacancy_assignments(db: Session, vacancy_id: int = None, recruiter_id: int = None):
    query = db.query(models.VacancyAssignment).options(
        joinedload(models.VacancyAssignment.recruiter),
        joinedload(models.VacancyAssignment.vacancy)
    )
    
    if vacancy_id:
        query = query.filter(models.VacancyAssignment.vacancy_id == vacancy_id)
    if recruiter_id:
        query = query.filter(models.VacancyAssignment.recruiter_id == recruiter_id)
    
    return query.order_by(models.VacancyAssignment.assigned_at.desc()).all()

def get_vacancy_assignment(db: Session, assignment_id: int):
    return db.query(models.VacancyAssignment).options(
        joinedload(models.VacancyAssignment.recruiter),
        joinedload(models.VacancyAssignment.vacancy)
    ).filter(models.VacancyAssignment.id == assignment_id).first()

def create_vacancy_assignment(db: Session, assignment: schemas.VacancyAssignmentCreate):
    # Проверяем, не назначена ли уже эта вакансия другому рекрутеру
    existing = db.query(models.VacancyAssignment).filter(
        models.VacancyAssignment.vacancy_id == assignment.vacancy_id,
        models.VacancyAssignment.status == "assigned"
    ).first()
    
    if existing:
        raise HTTPException(status_code=400, detail="Вакансия уже назначена другому рекрутеру")
    
    # Обновляем статус вакансии
    vacancy = get_vacancy(db, assignment.vacancy_id)
    vacancy.status = "in_progress"
    db.add(vacancy)
    
    # Создаем назначение
    db_assignment = models.VacancyAssignment(**assignment.model_dump())
    db.add(db_assignment)
    db.commit()
    db.refresh(db_assignment)
    
    # Загружаем информацию о рекрутере
    recruiter = db.query(models.User).filter(models.User.id == assignment.recruiter_id).first()
    db_assignment.recruiter = recruiter
    
    return db_assignment

def close_vacancy_assignment(db: Session, assignment_id: int, candidate_id: int):
    db_assignment = get_vacancy_assignment(db, assignment_id)
    if not db_assignment:
        return None
    
    # Обновляем назначение
    db_assignment.status = "closed"
    db_assignment.closed_at = datetime.utcnow()
    db_assignment.candidate_id = candidate_id
    
    # Обновляем статус вакансии
    vacancy = get_vacancy(db, db_assignment.vacancy_id)
    vacancy.status = "closed"
    
    db.add(db_assignment)
    db.add(vacancy)
    db.commit()
    db.refresh(db_assignment)
    return db_assignment

# Маршруты API для вакансий

@app.get("/vacancies/", response_model=List[schemas.VacancyDetail])
def read_vacancies(
    skip: int = 0, 
    limit: int = 100, 
    status: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    # Получаем список вакансий с подгрузкой связанных данных
    query = db.query(models.Vacancy).options(
        joinedload(models.Vacancy.creator),
        joinedload(models.Vacancy.assignments).joinedload(models.VacancyAssignment.recruiter)
    )
    
    if status:
        query = query.filter(models.Vacancy.status == status)
    
    vacancies = query.offset(skip).limit(limit).all()
    
    # Дополнительно загружаем рекрутеров для каждого назначения
    for vacancy in vacancies:
        for assignment in vacancy.assignments:
            if assignment.recruiter_id and not assignment.recruiter:
                # Явно загружаем информацию о рекрутере
                recruiter = db.query(models.User).filter(models.User.id == assignment.recruiter_id).first()
                if recruiter:
                    assignment.recruiter = recruiter
    
    return vacancies

@app.post("/vacancies/", response_model=schemas.Vacancy)
def create_new_vacancy(
    vacancy: schemas.VacancyCreate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    # Только администраторы могут создавать вакансии
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Только администраторы могут создавать вакансии")
    return create_vacancy(db, vacancy, current_user.id)

@app.get("/vacancies/{vacancy_id}", response_model=schemas.VacancyDetail)
def read_vacancy(
    vacancy_id: int = FastAPIPath(..., title="ID вакансии"),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    vacancy = get_vacancy(db, vacancy_id)
    if vacancy is None:
        raise HTTPException(status_code=404, detail="Вакансия не найдена")
    return vacancy

@app.put("/vacancies/{vacancy_id}", response_model=schemas.Vacancy)
def update_vacancy_endpoint(
    vacancy_id: int,
    vacancy: schemas.VacancyUpdate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    # Только администраторы могут обновлять вакансии
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Только администраторы могут обновлять вакансии")
    
    db_vacancy = update_vacancy(db, vacancy_id, vacancy)
    if db_vacancy is None:
        raise HTTPException(status_code=404, detail="Вакансия не найдена")
    return db_vacancy

@app.delete("/vacancies/{vacancy_id}", response_model=dict)
def delete_vacancy_endpoint(
    vacancy_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    # Только администраторы могут удалять вакансии
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Только администраторы могут удалять вакансии")
    
    result = delete_vacancy(db, vacancy_id)
    if not result:
        raise HTTPException(status_code=404, detail="Вакансия не найдена")
    return {"status": "success", "message": "Вакансия удалена"}

# Маршруты API для назначений вакансий

@app.post("/vacancy-assignments/", response_model=schemas.VacancyAssignmentDetail)
def create_assignment(
    assignment: schemas.VacancyAssignmentCreate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    # Проверяем, существует ли указанная вакансия
    vacancy = db.query(models.Vacancy).filter(models.Vacancy.id == assignment.vacancy_id).first()
    if not vacancy:
        raise HTTPException(status_code=404, detail="Вакансия не найдена")
    
    # Проверяем, существует ли указанный рекрутер
    recruiter = db.query(models.User).filter(
        models.User.id == assignment.recruiter_id,
        models.User.role == "recruiter"
    ).first()
    if not recruiter:
        raise HTTPException(status_code=404, detail="Рекрутер не найден")
    
    # Разные проверки для админов и рекрутеров
    if current_user.role == "admin":
        # Администраторы могут назначать любые вакансии любым рекрутерам
        pass
    elif current_user.role == "recruiter":
        # Рекрутеры могут принимать только свои вакансии
        if assignment.recruiter_id != current_user.id:
            raise HTTPException(status_code=403, detail="Вы можете назначать вакансии только себе")
    else:
        raise HTTPException(status_code=403, detail="Только администраторы и рекрутеры могут управлять назначениями вакансий")
    
    try:
        # Проверяем, не назначена ли уже эта вакансия другому рекрутеру
        existing_assignment = db.query(models.VacancyAssignment).filter(
            models.VacancyAssignment.vacancy_id == assignment.vacancy_id,
            models.VacancyAssignment.status == "assigned"
        ).first()
        
        if existing_assignment:
            # Если это переназначение вакансии администратором
            if current_user.role == "admin" and existing_assignment.recruiter_id != assignment.recruiter_id:
                # Отправляем уведомление предыдущему рекрутеру о переназначении
                previous_recruiter_id = existing_assignment.recruiter_id
                previous_recruiter = db.query(models.User).filter(models.User.id == previous_recruiter_id).first()
                
                notification = models.Notification(
                    user_id=previous_recruiter_id,
                    vacancy_id=assignment.vacancy_id,
                    message=f"Вакансия '{vacancy.title}' была переназначена другому рекрутеру.",
                    is_read=False
                )
                db.add(notification)
                
                # Закрываем старое назначение
                existing_assignment.status = "closed"
                existing_assignment.closed_at = datetime.now()
                db.commit()
            else:
                raise HTTPException(status_code=400, detail="Эта вакансия уже назначена рекрутеру")
        
        # Создаем назначение и загружаем связанные данные
        db_assignment = create_vacancy_assignment(db, assignment)
        
        # Отправляем уведомление рекрутеру о назначении
        notification = models.Notification(
            user_id=assignment.recruiter_id,
            vacancy_id=assignment.vacancy_id,
            message=f"Вам назначена новая вакансия: {vacancy.title}",
            is_read=False
        )
        db.add(notification)
        db.commit()
        
        # Загружаем связанные данные для возврата полной информации
        db_assignment_full = db.query(models.VacancyAssignment).options(
            joinedload(models.VacancyAssignment.recruiter),
            joinedload(models.VacancyAssignment.vacancy)
        ).filter(models.VacancyAssignment.id == db_assignment.id).first()
        
        return db_assignment_full
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.put("/vacancy-assignments/{assignment_id}/close", response_model=schemas.VacancyAssignment)
def close_assignment(
    assignment_id: int,
    candidate_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    # Проверяем, существует ли назначение
    assignment = get_vacancy_assignment(db, assignment_id)
    if not assignment:
        raise HTTPException(status_code=404, detail="Назначение не найдено")
    
    # Проверяем, принадлежит ли назначение текущему пользователю
    if assignment.recruiter_id != current_user.id:
        raise HTTPException(status_code=403, detail="Вы можете закрывать только свои назначения")
    
    # Проверяем, не закрыто ли уже назначение
    if assignment.status == "closed":
        raise HTTPException(status_code=400, detail="Назначение уже закрыто")
    
    # Закрываем назначение
    return close_vacancy_assignment(db, assignment_id, candidate_id)

@app.get("/vacancy-assignments/", response_model=List[schemas.VacancyAssignment])
def read_assignments(
    vacancy_id: Optional[int] = None,
    recruiter_id: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    # Если пользователь - рекрутер, он может видеть только свои назначения
    if current_user.role == "recruiter":
        recruiter_id = current_user.id
    
    return get_vacancy_assignments(db, vacancy_id, recruiter_id)

@app.get("/vacancy-assignments/{assignment_id}", response_model=schemas.VacancyAssignmentDetail)
def read_assignment(
    assignment_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    assignment = get_vacancy_assignment(db, assignment_id)
    if not assignment:
        raise HTTPException(status_code=404, detail="Назначение не найдено")
    
    # Если пользователь - рекрутер, он может видеть только свои назначения
    if current_user.role == "recruiter" and assignment.recruiter_id != current_user.id:
        raise HTTPException(status_code=403, detail="У вас нет доступа к этому назначению")
    
    return assignment

@app.put("/vacancy-assignments/{assignment_id}/close", response_model=schemas.VacancyAssignment)
def close_vacancy_assignment(
    assignment_id: int = FastAPIPath(..., title="ID назначения"),
    candidate_id: Optional[int] = Query(None, title="ID кандидата"),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """
    Закрыть назначение вакансии (рекрутер нашел кандидата).
    """
    assignment = db.query(models.VacancyAssignment).filter(
        models.VacancyAssignment.id == assignment_id
    ).first()
    
    if not assignment:
        raise HTTPException(status_code=404, detail="Назначение не найдено")
    
    # Проверяем, что пользователь - это назначенный рекрутер или администратор
    if current_user.id != assignment.recruiter_id and current_user.role != "admin":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="У вас нет прав для закрытия этого назначения"
        )
    
    # Проверяем, что назначение еще не закрыто
    if assignment.status == "closed":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Это назначение уже закрыто"
        )
    
    # Обновляем назначение
    assignment.status = "closed"
    assignment.closed_at = datetime.now()
    assignment.candidate_id = candidate_id
    
    # Также закрываем связанную вакансию
    vacancy = db.query(models.Vacancy).filter(
        models.Vacancy.id == assignment.vacancy_id
    ).first()
    
    if vacancy:
        vacancy.status = "closed"
        vacancy.updated_at = datetime.now()
    
    db.commit()
    db.refresh(assignment)
    
    # Отправляем уведомление администраторам
    admins = db.query(models.User).filter(models.User.role == "admin").all()
    recruiter = db.query(models.User).filter(models.User.id == assignment.recruiter_id).first()
    
    for admin in admins:
        message = f"Рекрутер {recruiter.full_name} закрыл вакансию: {vacancy.title}"
        if candidate_id:
            candidate = db.query(models.Application).filter(models.Application.id == candidate_id).first()
            if candidate:
                message += f" с кандидатом {candidate.name}"
        
        notification = models.Notification(
            user_id=admin.id,
            vacancy_id=vacancy.id,
            message=message,
            is_read=False
        )
        db.add(notification)
    
    db.commit()
    
    return assignment

@app.put("/vacancies/{vacancy_id}", response_model=schemas.Vacancy)
def update_vacancy(
    vacancy_data: schemas.VacancyUpdate,
    vacancy_id: int = FastAPIPath(..., title="ID вакансии"),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    # Только администраторы могут обновлять вакансии
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Только администраторы могут обновлять вакансии")
    
    db_vacancy = update_vacancy(db, vacancy_id, vacancy_data)
    if db_vacancy is None:
        raise HTTPException(status_code=404, detail="Вакансия не найдена")
    return db_vacancy

@app.put("/vacancies/{vacancy_id}/close", response_model=schemas.Vacancy)
def close_vacancy_endpoint(
    vacancy_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """
    Close a vacancy directly without requiring a candidate.
    """
    # Check if vacancy exists
    vacancy = db.query(models.Vacancy).filter(models.Vacancy.id == vacancy_id).first()
    if not vacancy:
        raise HTTPException(status_code=404, detail="Вакансия не найдена")
    
    # Check if user is admin or recruiter assigned to this vacancy
    is_assigned = False
    if current_user.role == "recruiter":
        # Check if this recruiter is assigned to the vacancy
        assignment = db.query(models.VacancyAssignment).filter(
            models.VacancyAssignment.vacancy_id == vacancy_id,
            models.VacancyAssignment.recruiter_id == current_user.id,
            models.VacancyAssignment.status == "assigned"
        ).first()
        
        if assignment:
            is_assigned = True
    
    if current_user.role != "admin" and not is_assigned:
        raise HTTPException(status_code=403, detail="У вас нет прав для закрытия этой вакансии")
    
    # Update vacancy status and closed information
    vacancy.status = "closed"
    vacancy.updated_at = datetime.now()
    vacancy.closed_by_id = current_user.id
    vacancy.closed_at = datetime.now()
    
    # If there's an assignment, close it too
    if is_assigned:
        assignment.status = "closed"
        assignment.closed_at = datetime.now()
    
    # Create notification for admins
    admins = db.query(models.User).filter(models.User.role == "admin").all()
    for admin in admins:
        if admin.id != current_user.id:  # Don't notify the admin who closed it
            notification = models.Notification(
                user_id=admin.id,
                vacancy_id=vacancy.id,
                message=f"{current_user.full_name} закрыл(а) вакансию: {vacancy.title}",
                is_read=False
            )
            db.add(notification)
    
    db.commit()
    db.refresh(vacancy)
    
    return vacancy

# Константы для путей хранения файлов
DOCUMENTS_DIR = BASE_DIR / "documents"
DOCUMENTS_DIR.mkdir(exist_ok=True)

# Document endpoints
@app.post("/folders/", response_model=schemas.Folder)
def create_folder(
    folder: schemas.FolderCreate,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    """Создание новой папки для текущего пользователя"""
    # Проверяем, существует ли родительская папка, если она указана
    if folder.parent_id:
        parent_folder = db.query(models.Folder).filter(
            models.Folder.id == folder.parent_id,
            models.Folder.owner_id == current_user.id
        ).first()
        if not parent_folder:
            raise HTTPException(status_code=404, detail="Parent folder not found or access denied")
    
    # Проверяем, существует ли уже папка с таким именем на том же уровне
    existing_folder = db.query(models.Folder).filter(
        models.Folder.name == folder.name,
        models.Folder.parent_id == folder.parent_id,
        models.Folder.owner_id == current_user.id
    ).first()
    
    if existing_folder:
        raise HTTPException(status_code=400, detail="Folder with this name already exists")
    
    # Создаем новую папку
    db_folder = models.Folder(
        **folder.dict(),
        owner_id=current_user.id
    )
    db.add(db_folder)
    db.commit()
    db.refresh(db_folder)
    return db_folder

@app.get("/folders/", response_model=List[schemas.Folder])
def read_folders(
    parent_id: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    """Получение списка папок пользователя на определенном уровне иерархии"""
    query = db.query(models.Folder).filter(models.Folder.owner_id == current_user.id)
    
    if parent_id is not None:
        query = query.filter(models.Folder.parent_id == parent_id)
    else:
        query = query.filter(models.Folder.parent_id == None)
    
    folders = query.all()
    return folders

@app.get("/folders/{folder_id}", response_model=schemas.FolderDetail)
def read_folder(
    folder_id: int,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    """Получение информации о конкретной папке вместе с вложенными папками и документами"""
    folder = db.query(models.Folder).filter(
        models.Folder.id == folder_id,
        models.Folder.owner_id == current_user.id
    ).first()
    
    if not folder:
        raise HTTPException(status_code=404, detail="Folder not found or access denied")
    
    return folder

@app.put("/folders/{folder_id}", response_model=schemas.Folder)
def update_folder(
    folder_id: int,
    folder_update: schemas.FolderUpdate,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    """Обновление информации о папке"""
    db_folder = db.query(models.Folder).filter(
        models.Folder.id == folder_id,
        models.Folder.owner_id == current_user.id
    ).first()
    
    if not db_folder:
        raise HTTPException(status_code=404, detail="Folder not found or access denied")
    
    # Проверяем, что не создаем циклическую ссылку
    if folder_update.parent_id and folder_update.parent_id == folder_id:
        raise HTTPException(status_code=400, detail="Cannot set folder as its own parent")
    
    # Проверяем, что родительская папка существует и принадлежит пользователю
    if folder_update.parent_id:
        parent_folder = db.query(models.Folder).filter(
            models.Folder.id == folder_update.parent_id,
            models.Folder.owner_id == current_user.id
        ).first()
        
        if not parent_folder:
            raise HTTPException(status_code=404, detail="Parent folder not found or access denied")
    
    # Обновляем поля папки
    if folder_update.name is not None:
        db_folder.name = folder_update.name
    
    if folder_update.parent_id is not None:
        db_folder.parent_id = folder_update.parent_id
    
    db.commit()
    db.refresh(db_folder)
    return db_folder

@app.delete("/folders/{folder_id}")
def delete_folder(
    folder_id: int,
    recursive: bool = False,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    """Удаление папки. Если recursive=True, удаляются все вложенные папки и документы"""
    db_folder = db.query(models.Folder).filter(
        models.Folder.id == folder_id,
        models.Folder.owner_id == current_user.id
    ).first()
    
    if not db_folder:
        raise HTTPException(status_code=404, detail="Folder not found or access denied")
    
    # Проверяем, есть ли вложенные папки или документы
    subfolders = db.query(models.Folder).filter(models.Folder.parent_id == folder_id).count()
    documents = db.query(models.Document).filter(models.Document.folder_id == folder_id).count()
    
    if (subfolders > 0 or documents > 0) and not recursive:
        raise HTTPException(
            status_code=400, 
            detail="Folder contains subfolders or documents. Set recursive=true to delete everything"
        )
    
    if recursive:
        # Удаляем все документы в папке и в подпапках рекурсивно
        def delete_folder_contents(folder_id):
            # Получаем все подпапки
            subfolders = db.query(models.Folder).filter(models.Folder.parent_id == folder_id).all()
            
            for subfolder in subfolders:
                delete_folder_contents(subfolder.id)
            
            # Удаляем все документы в папке
            documents = db.query(models.Document).filter(models.Document.folder_id == folder_id).all()
            
            for document in documents:
                # Удаляем физический файл
                file_path = DOCUMENTS_DIR / document.file_path
                if file_path.exists():
                    file_path.unlink()
                
                # Удаляем записи о доступе
                db.query(models.DocumentAccess).filter(models.DocumentAccess.document_id == document.id).delete()
            
            # Удаляем все документы в папке
            db.query(models.Document).filter(models.Document.folder_id == folder_id).delete()
            
            # Удаляем подпапки
            db.query(models.Folder).filter(models.Folder.parent_id == folder_id).delete()
        
        delete_folder_contents(folder_id)
    
    # Удаляем саму папку
    db.delete(db_folder)
    db.commit()
    
    return {"status": "success", "message": "Folder deleted"}

@app.post("/documents/", response_model=schemas.Document)
async def create_document(
    folder_id: Optional[int] = Form(None),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    """Загрузка нового документа"""
    # Проверяем, существует ли указанная папка и принадлежит ли она пользователю
    if folder_id:
        folder = db.query(models.Folder).filter(
            models.Folder.id == folder_id,
            models.Folder.owner_id == current_user.id
        ).first()
        
        if not folder:
            raise HTTPException(status_code=404, detail="Folder not found or access denied")
    
    # Проверяем, есть ли уже документ с таким именем в этой папке
    existing_document = db.query(models.Document).filter(
        models.Document.name == file.filename,
        models.Document.folder_id == folder_id,
        models.Document.owner_id == current_user.id
    ).first()
    
    if existing_document:
        raise HTTPException(status_code=400, detail="Document with this name already exists in this folder")
    
    # Генерируем уникальное имя файла
    file_extension = file.filename.split(".")[-1] if "." in file.filename else ""
    unique_filename = f"{uuid4()}.{file_extension}" if file_extension else f"{uuid4()}"
    
    # Создаем директорию для пользователя, если она не существует
    user_dir = DOCUMENTS_DIR / str(current_user.id)
    user_dir.mkdir(exist_ok=True)
    
    # Путь для сохранения файла
    file_path = user_dir / unique_filename
    
    # Сохраняем файл
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())
    
    # Получаем размер файла
    file_size = os.path.getsize(file_path)
    
    # Определяем тип файла
    content_type = file.content_type or "application/octet-stream"
    
    # Создаем запись в базе данных
    db_document = models.Document(
        name=file.filename,
        file_path=str(int(current_user.id)) + "/" + unique_filename,
        file_type=content_type,
        file_size=file_size,
        folder_id=folder_id,
        owner_id=current_user.id,
        created_at=datetime.now(),
        updated_at=datetime.now()
    )
    
    db.add(db_document)
    db.commit()
    db.refresh(db_document)
    
    return db_document

@app.get("/documents/", response_model=List[schemas.Document])
def read_documents(
    folder_id: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    """Получение списка документов пользователя в указанной папке"""
    query = db.query(models.Document).filter(models.Document.owner_id == current_user.id)
    
    if folder_id is not None:
        query = query.filter(models.Document.folder_id == folder_id)
    
    documents = query.all()
    
    # Проверяем и фиксируем NULL значения в полях created_at и updated_at
    current_time = datetime.now()
    for doc in documents:
        if doc.created_at is None:
            doc.created_at = current_time
            db.add(doc)
        if doc.updated_at is None:
            doc.updated_at = current_time
            db.add(doc)
    
    # Сохраняем изменения, если были исправления
    db.commit()
    
    return documents

@app.get("/documents/{document_id}", response_model=schemas.Document)
def read_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    """Получение информации о конкретном документе"""
    document = db.query(models.Document).filter(
        models.Document.id == document_id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Проверяем, имеет ли пользователь доступ к документу
    if document.owner_id != current_user.id:
        # Проверяем, есть ли у пользователя доступ через таблицу document_access
        access = db.query(models.DocumentAccess).filter(
            models.DocumentAccess.document_id == document_id,
            models.DocumentAccess.user_id == current_user.id
        ).first()
        
        if not access:
            raise HTTPException(status_code=403, detail="Access denied")
    
    # Проверяем и фиксируем NULL значения в полях created_at и updated_at
    current_time = datetime.now()
    if document.created_at is None:
        document.created_at = current_time
        db.add(document)
    if document.updated_at is None:
        document.updated_at = current_time
        db.add(document)
    
    # Сохраняем изменения, если были исправления
    db.commit()
    
    return document

@app.get("/documents/{document_id}/download")
async def download_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    """Скачивание документа"""
    document = db.query(models.Document).filter(
        models.Document.id == document_id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Проверяем, имеет ли пользователь доступ к документу
    if document.owner_id != current_user.id:
        # Проверяем, есть ли у пользователя доступ через таблицу document_access
        access = db.query(models.DocumentAccess).filter(
            models.DocumentAccess.document_id == document_id,
            models.DocumentAccess.user_id == current_user.id
        ).first()
        
        if not access:
            raise HTTPException(status_code=403, detail="Access denied")
    
    # Проверяем и фиксируем NULL значения в полях created_at и updated_at
    current_time = datetime.now()
    if document.created_at is None:
        document.created_at = current_time
        db.add(document)
    if document.updated_at is None:
        document.updated_at = current_time
        db.add(document)
    
    # Сохраняем изменения, если были исправления
    db.commit()
    
    file_path = DOCUMENTS_DIR / document.file_path
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        path=file_path,
        filename=document.name,
        media_type=document.file_type
    )

@app.put("/documents/{document_id}", response_model=schemas.Document)
def update_document(
    document_id: int,
    document_update: schemas.DocumentUpdate,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    """Обновление информации о документе (имя или папка)"""
    document = db.query(models.Document).filter(
        models.Document.id == document_id,
        models.Document.owner_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found or access denied")
    
    # Проверяем, существует ли указанная папка и принадлежит ли она пользователю
    if document_update.folder_id is not None:
        if document_update.folder_id != 0:  # 0 означает корневую папку (None)
            folder = db.query(models.Folder).filter(
                models.Folder.id == document_update.folder_id,
                models.Folder.owner_id == current_user.id
            ).first()
            
            if not folder:
                raise HTTPException(status_code=404, detail="Folder not found or access denied")
            
            document.folder_id = document_update.folder_id
        else:
            document.folder_id = None
    
    if document_update.name is not None:
        # Проверяем, есть ли уже документ с таким именем в этой папке
        existing_document = db.query(models.Document).filter(
            models.Document.name == document_update.name,
            models.Document.folder_id == document.folder_id,
            models.Document.owner_id == current_user.id,
            models.Document.id != document_id
        ).first()
        
        if existing_document:
            raise HTTPException(
                status_code=400, 
                detail="Document with this name already exists in this folder"
            )
        
        document.name = document_update.name
    
    # Проверяем и фиксируем NULL значения в полях created_at и updated_at
    current_time = datetime.now()
    if document.created_at is None:
        document.created_at = current_time
    
    # Всегда обновляем updated_at при изменении документа
    document.updated_at = current_time
    
    db.commit()
    db.refresh(document)
    return document

@app.delete("/documents/{document_id}")
def delete_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    """Удаление документа"""
    document = db.query(models.Document).filter(
        models.Document.id == document_id,
        models.Document.owner_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found or access denied")
    
    # Удаляем физический файл
    file_path = DOCUMENTS_DIR / document.file_path
    if file_path.exists():
        file_path.unlink()
    
    # Удаляем записи о доступе
    db.query(models.DocumentAccess).filter(models.DocumentAccess.document_id == document_id).delete()
    
    # Удаляем документ из базы данных
    db.delete(document)
    db.commit()
    
    return {"status": "success", "message": "Document deleted"}

@app.post("/documents/{document_id}/share", response_model=schemas.DocumentAccess)
def share_document(
    document_id: int,
    access_data: schemas.DocumentAccessCreate,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    """Предоставление доступа к документу другому пользователю"""
    # Проверяем, существует ли документ и принадлежит ли он текущему пользователю
    document = db.query(models.Document).filter(
        models.Document.id == document_id,
        models.Document.owner_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found or access denied")
    
    # Проверяем, существует ли пользователь, которому предоставляется доступ
    user = db.query(models.User).filter(models.User.id == access_data.user_id).first()
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Проверяем, нет ли уже доступа
    existing_access = db.query(models.DocumentAccess).filter(
        models.DocumentAccess.document_id == document_id,
        models.DocumentAccess.user_id == access_data.user_id
    ).first()
    
    if existing_access:
        # Обновляем уровень доступа
        existing_access.access_level = access_data.access_level
        db.commit()
        db.refresh(existing_access)
        return existing_access
    
    # Создаем запись о доступе
    db_access = models.DocumentAccess(
        document_id=document_id,
        user_id=access_data.user_id,
        access_level=access_data.access_level
    )
    
    db.add(db_access)
    db.commit()
    db.refresh(db_access)
    
    return db_access

@app.delete("/documents/{document_id}/share/{user_id}")
def revoke_document_access(
    document_id: int,
    user_id: int,
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user)
):
    """Отзыв доступа к документу у пользователя"""
    # Проверяем, существует ли документ и принадлежит ли он текущему пользователю
    document = db.query(models.Document).filter(
        models.Document.id == document_id,
        models.Document.owner_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found or access denied")
    
    # Удаляем запись о доступе
    result = db.query(models.DocumentAccess).filter(
        models.DocumentAccess.document_id == document_id,
        models.DocumentAccess.user_id == user_id
    ).delete()
    
    if result == 0:
        raise HTTPException(status_code=404, detail="Access record not found")
    
    db.commit()
    
    return {"status": "success", "message": "Access revoked"}

@app.post("/telegram/broadcast", response_model=Dict[str, Any])
async def send_telegram_broadcast(
    broadcast: schemas.TelegramBroadcast,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """
    Отправляет сообщение в Telegram пользователям.
    Можно выбрать получателей по статусу или индивидуально по ID.
    Сообщения будут содержать кнопки подтверждения получения.
    """
    # Проверяем, что пользователь является администратором
    if current_user.role != "admin":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Только администраторы могут отправлять массовые рассылки"
        )
    
    # Импортируем сервис telegram асинхронно, чтобы избежать циклических импортов
    from telegram_service import process_broadcast
    
    try:
        # Создаем запись о рассылке в БД
        broadcast_record = models.TelegramBroadcast(
            sender_id=current_user.id,
            message=broadcast.message,
            status="processing"
        )
        db.add(broadcast_record)
        db.commit()
        db.refresh(broadcast_record)
        
        # Обрабатываем рассылку с передачей ID рассылки
        result = await process_broadcast(
            db=db,
            message=broadcast.message,
            status=broadcast.status,
            user_ids=broadcast.user_ids,
            broadcast_id=broadcast_record.id  # Передаем ID рассылки
        )
        
        # Обновляем запись о рассылке
        broadcast_record.status = "completed"
        broadcast_record.recipients_count = result.get("results", {}).get("success_count", 0)
        db.commit()
        
        return result
    except Exception as e:
        # В случае ошибки обновляем статус рассылки
        if 'broadcast_record' in locals():
            broadcast_record.status = "failed"
            db.commit()
        
        logging.error(f"Error sending Telegram broadcast: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Ошибка при отправке сообщений: {str(e)}"
        )

# Добавить после эндпоинта для отправки сообщений

@app.get("/telegram/broadcasts", response_model=List[schemas.TelegramBroadcastResponse])
def get_broadcast_history(
    skip: int = 0,
    limit: int = 20,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """
    Получает историю рассылок сообщений в Telegram.
    Только для администраторов.
    """
    # Проверяем, что пользователь является администратором
    if current_user.role != "admin":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Только администраторы могут просматривать историю рассылок"
        )
    
    # Получаем историю рассылок
    broadcasts = db.query(models.TelegramBroadcast)\
        .order_by(models.TelegramBroadcast.created_at.desc())\
        .offset(skip)\
        .limit(limit)\
        .all()
    
    return broadcasts

@app.get("/telegram/broadcasts/{broadcast_id}", response_model=schemas.TelegramBroadcastDetailResponse)
async def get_broadcast_details(
    broadcast_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """
    Получает детальную информацию о рассылке, включая подтверждения получения.
    """
    if current_user.role != "admin":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Только администраторы могут получать детали рассылок"
        )
    
    broadcast = db.query(models.TelegramBroadcast).filter(
        models.TelegramBroadcast.id == broadcast_id
    ).first()
    
    if not broadcast:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Рассылка с ID {broadcast_id} не найдена"
        )
    
    return broadcast

@app.get("/telegram/broadcasts/{broadcast_id}/confirmations", response_model=List[schemas.MessageConfirmationResponse])
async def get_broadcast_confirmations(
    broadcast_id: int,
    status: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """
    Получает список подтверждений для конкретной рассылки с возможностью фильтрации.
    """
    if current_user.role != "admin":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Только администраторы могут получать детали рассылок"
        )
    
    query = db.query(models.MessageConfirmation).filter(
        models.MessageConfirmation.broadcast_id == broadcast_id
    )
    
    if status:
        query = query.filter(models.MessageConfirmation.status == status)
    
    confirmations = query.all()
    
    if not confirmations:
        # Возвращаем пустой список, а не ошибку
        return []
    
    return confirmations